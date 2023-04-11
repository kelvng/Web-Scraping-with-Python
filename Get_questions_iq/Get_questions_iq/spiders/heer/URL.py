import json
import sys
from hashlib import new
import csv
import requests
from bs4 import BeautifulSoup
from docx import Document
from scrapy.selector import Selector


def url_obj(url, param, param1):
    pass


# Tên spider
name = "Geturl"

# resp = requests.post("https://os.3i.com.vn/PythonCrawler/GetCrawlerData?spiderName=crawler")
# resp_json = resp.json()
# # load post tạo json
# print(resp_json)
# print(resp_json['Url'])
# url = resp_json['Url']
url = "https://suretest.vn/"

date = []
Title = []
author = []
image = []
summary = []
list_child_url = []
content = []


# mảng chứa lisr url và content text

class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep

    Idx = 0
    chk = 0

    # Hàm kiểm tra link đã tồn tại trong list chưa
    def chk_link_exist(self, link):
        for obj in list_child_url:
            if (obj.url == link):
                return 1
        return 0

    # Hàm quét và lấy các link con từ link mẹ
    def Extract_Url(self, url, deep):
        try:
            # Độ sâu của link mẹ là 1, link con = link mẹ +1
            if deep <= 6:
                self.deep = deep
                req = requests.get(url)
                html = req.text
                soup = BeautifulSoup(html, 'html.parser')
                LstLink = soup.find('body')
                for s in LstLink.find_all('a'):
                    try:
                        link = s['href']
                        if 'http' not in link:
                            link = 'https://suretest.vns' + link
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        else:
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                    except:
                        pass
        except:
            print('link fail')
            pass

    # with open('C:\\Users\\pycha\\PycharmProjects\\Source\\DomainManagement\\DomainManagement\\spiders\\selector2.json',
    #           'r') as j:
    #     list_tag = json.loads(j.read())

    # def GetContent(self, url):
    #     list_tag = self.list_tag
    #     req = requests.get(url)
    #     html = req.text
    #     soup = BeautifulSoup(html, 'html.parser')
    #     LstLink = soup.find('html')
    #     for s in LstLink.find_all(list_tag['title']):
    #         title = s.text
    #         if '' != title:
    #             content.append(title)
    #
    #     return content

    def main(self):
        Idx = 0
        len_list = len(list_child_url)
        while Idx < len_list:
            self.Extract_Url(list_child_url[Idx].url, list_child_url[Idx].deep)
            Idx = Idx + 1
            len_list = len(list_child_url)


list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
object.main()
document = Document()
# object.GetContent(url)
list_url = []
list_deep = []
list_iscan = []
for obj in list_child_url:
    # list_url.append(obj.url)
    # list_deep.append(obj.deep)
    # list_iscan.append(obj.iscan)
    # document.add_paragraph(obj.url)
    print(obj.url, obj.iscan, obj.deep, sep=' ')
print(len(list_child_url))
for value in content:
    document.add_paragraph(value)
    print(value)
# for sum in summary:
#     document.add_paragraph(sum)
#     print(sum)
print(len(content))

# Vị trí lưu file

url = url.replace("https://", "")
filename = url.replace(".vn/","") + "deep=4"+".docx"
filePath = r'\Get_questions_iq\Get_questions_iq/' + "url.docx"
document.save(r"C:\Users\Admin\PycharmProjects\Source" + filePath)
headerList = ['link','iscan', 'deep']
# rows = zip(list_url,list_iscan,list_deep)
# with open(r"C:\Users\Admin\PycharmProjects\Source" + '789.csv', 'w',newline= '') as file:
#     writer = csv.writer(file )
#     writer.writerow(headerList)
#     for row in rows:
#         writer.writerow(row)

# #up_load file lên server:
# url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
# #url_upload = resp_json['DataStoragePath']
# response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "phancuoc_cntt_3i"}, files={
#     "fileUpload": (
#     'url.docx', open("E:\Source" + filePath, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
# if response_upload.ok:
#     print("Upload completed successfully!")
#     print(response_upload.text)
# else:
#     print("Something went wrong!")
#get link deep 5 hocmai
exit()