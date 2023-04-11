# coding=utf8
# -*- coding: utf-8 -*-
import csv
import json
import os
import re
from datetime import datetime

import numpy as np
from pywinauto import Application
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import requests
from bs4 import BeautifulSoup
import scrapy
from time import sleep
import pandas as pd
from bs4 import BeautifulSoup as BSHTML
import uuid
from tqdm import tqdm
from docx import Document
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC


options = webdriver.ChromeOptions()
prefs = {"profile.default_content_setting_values.notifications": 2}

options.add_experimental_option("prefs", prefs)
options.add_argument("--start-maximized")
options.add_argument("log-level=3")
# options.add_argument('--headless')
browser = webdriver.Chrome(r'F:\PycharmProjects\Source\chromedriver.exe', options=options)
browser.get('https://suretest.vn/lam-bai-thi/de-thi-toan-thpt-quoc-gia-nam-2019-ma-de-121-bo-giao-duc-va-dao-tao-143284.html')
sleep(2)
#login
mail = browser.find_element_by_xpath('//*[@id="username"]')
mail.send_keys('zetabase3i@gmail.com')
password = browser.find_element_by_xpath('//*[@id="password"]')
password.send_keys('Langnghiem79')
browser.find_element_by_xpath('//*[@class="btn btn-success key-enter"]').click()
list_child_url = []
url = 'https://suretest.vn'

class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep

    Idx = 0
    chk = 0

    # Hàm kiểm tra link đã tồn tại trong list chưa
    def chk_link_exist(self, link):
        for obj in list_child_url:
            if (obj.url == link):
                return 1
        return 0

    # Hàm quét và lấy các link con từ link mẹ
    def Extract_Url(self, url, deep):
        try:
            # Độ sâu của link mẹ là 1, link con = link mẹ +1
            if deep <= 6:
                self.deep = deep
                browser.get(url)
                wait = WebDriverWait(browser, 10)
                email = wait.until(EC.visibility_of_element_located((By.TAG_NAME, 'body')))
                page_source = browser.page_source
                soup = BeautifulSoup(page_source, 'html.parser')
                print(url)
                #LstLink = soup.find('body')

                for s in soup.find_all('a'):
                    try:
                        link = s['href']
                        if 'http' not in link:
                            link = 'https://suretest.vn' + link
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        # elif 'http' not in link and '#' in link:
                        #     print('tier 2 link')
                        #     pass
                        else:
                            if 'suretest.vn' in link:
                                if self.chk_link_exist(link) == 0:
                                    #if 'https://violet.vn' in link:
                                    url = link
                                    iscan = 0
                                    deep = self.deep + 1
                                    list_child_url.append(url_obj(url, iscan, deep))
                            else:
                                print('strange link!')
                                pass
                    except:
                        pass
        except:
            print('link fail')
            pass

    def main(self):
        Idx = 0
        len_list = len(list_child_url)
        while Idx < len_list:
            self.Extract_Url(list_child_url[Idx].url, list_child_url[Idx].deep)
            Idx = Idx + 1
            len_list = len(list_child_url)


list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
object.main()
document = Document()
deep = Document()
# object.GetContent(url)
for obj in list_child_url:
    document.add_paragraph(obj.url)
    print(obj.url, obj.iscan, obj.deep, sep=' ')
print(len(list_child_url))
for obj in list_child_url:
    document.add_paragraph(str(obj.deep))

# for sum in summary:
#     document.add_paragraph(sum)
#     print(sum)


# Vị trí lưu file
print('save')
filename = 'suretest268'

with open(filename+'.csv', 'w',  newline = '',encoding="utf-8") as file_output:
    headers = ['link']
    writer = csv.DictWriter(file_output, delimiter=',', lineterminator='\n',fieldnames=headers)
    writer.writeheader()
    #print(URLs_all_page)
    for value in list_child_url:
        print(value.url)
        writer.writerow({headers[0]: value.url})
filename = 'suretest268' + '.docx'
print(filename)
filePath = r'C:\Users\Admin 3i\PycharmProjects\Source\Get_questions_iq\Get_questions_iq\spiders\tracnghiemnet/'+ filename
document.save(filename)
exit()