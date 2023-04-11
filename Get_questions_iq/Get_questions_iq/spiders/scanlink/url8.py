# coding=utf8
# -*- coding: utf-8 -*-
import csv
import json
import sys
from hashlib import new
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import requests
from bs4 import BeautifulSoup
from docx import Document
from scrapy.selector import Selector
from scrapy.spiders import CrawlSpider, Rule

from scrapy.linkextractors import LinkExtractor
def url_obj(url, param, param1):
    pass


# Tên spider
name = "Geturl"

# resp = requests.post("https://os.3i.com.vn/PythonCrawler/GetCrawlerData?spiderName=crawler")
# resp_json = resp.json()
# # load post tạo json
# print(resp_json)
# print(resp_json['Url'])
# url = resp_json['Url']
url = "https://tracnghiem.net/"


date = []
Title = []
author = []
image = []
summary = []
list_child_url = []
content = []


# mảng chứa lisr url và content text
rules = (
        Rule(LinkExtractor(allow=('')), callback="parse_page", follow=True),
    )
class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep

    Idx = 0
    chk = 0

    # Hàm kiểm tra link đã tồn tại trong list chưa
    def chk_link_exist(self, link):
        for obj in list_child_url:
            if (obj.url == link):
                return 1
        return 0

    # Hàm quét và lấy các link con từ link mẹ
    def Extract_Url(self, url, deep):
        try:
            # Độ sâu của link mẹ là 1, link con = link mẹ +1
            if deep <= 8:
                self.deep = deep
                print(url)
                # payload = {'txtLoginUsername': 'Langnghiem79', 'txtLoginPassword': 'Langnghiem79'}
                # login = 'https://tailieu.vn/'
                #
                # with requests.Session() as session:
                #     post = session.post(login, data=payload,verify=False)
                #     req = session.get(url, verify=False)
                req = requests.get(url, verify=False)
                html = req.text
                soup = BeautifulSoup(html, 'html.parser')
                LstLink = soup.find('body')

                for s in LstLink.find_all('a'):
                    try:
                        link = s['href']
                        if 'http' not in link:
                            link = 'https://tracnghiem.net' + link
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        else:
                            if 'tracnghiem.net' in link:

                                if self.chk_link_exist(link) == 0:
                                   # if 'vatly247.com' in link:
                                        url = link
                                        iscan = 0
                                        deep = self.deep + 1
                                        list_child_url.append(url_obj(url, iscan, deep))
                            else:
                                print('strange link!')
                                pass
                    except:
                        pass
        except:
            print('link fail')
            pass

    def main(self):
        Idx = 0
        len_list = len(list_child_url)
        while Idx < len_list:
            self.Extract_Url(list_child_url[Idx].url, list_child_url[Idx].deep)
            Idx = Idx + 1
            len_list = len(list_child_url)


list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
object.main()
document = Document()
deep = Document()
# object.GetContent(url)
for obj in list_child_url:
    document.add_paragraph(obj.url)
    print(obj.url, obj.iscan, obj.deep, sep=' ')
print(len(list_child_url))
for obj in list_child_url:
    document.add_paragraph(str(obj.deep))
for value in content:
    document.add_paragraph(value)
    #print(value)
# for sum in summary:
#     document.add_paragraph(sum)
#     print(sum)
print(len(content))
#https://sachgiaibaitap.com
filename = 'tracnghiem'
with open('tracnghiem.csv', 'w',  newline = '',encoding="utf-8") as file_output:
    headers = ['link']
    writer = csv.DictWriter(file_output, delimiter=',', lineterminator='\n',fieldnames=headers)
    writer.writeheader()
    #print(URLs_all_page)
    for value in list_child_url:
        print(value.url)
        writer.writerow({headers[0]: value.url})

# Vị trí lưu file
print('save')
filename = 'tracnghiem' + '.docx'
print(filename)
filePath = r'H:\PycharmProjects\Source\Get_questions_iq\Get_questions_iq\spiders\scanlink/'+ filename
document.save(filename)
exit()