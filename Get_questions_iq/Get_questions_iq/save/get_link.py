import json
import requests
from bs4 import BeautifulSoup
from docx import Document


def url_obj(url, param, param1):
    pass
#Tên spider
name ="Geturl"


resp = requests.post("https://os.3i.com.vn/PythonCrawler/GetCrawlerData?spiderName=crawler")
resp_json = resp.json()
print(resp_json)

url = resp_json['Url']

list_child_url = []
content = []


class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep
    Idx = 0;
    chk = 0;
    # Hàm kiểm tra link đã tồn tại trong list chưa
    def chk_link_exist(self, link):
        for obj in list_child_url:
            if (obj.url == link):
                 return 1;
        return 0;
    #Hàm quét và lấy các link con từ link mẹ
    def Extract_Url(self,url,deep):
        try:
            # Độ sâu của link mẹ là 1, link con = link mẹ +1
            if deep <= 1:
                req = requests.get(url)
                html = req.text
                soup = BeautifulSoup(html, 'html.parser')
                LstLink = soup.find('body')
                for s in LstLink.find_all('a'):
                    try:
                        link = s['href']
                        if 'https' not in link:
                            link = 'https://vneconomy.vn' + link
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        else:
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                    except:
                        pass
        except:
            print('link fail')
            pass
    with open('E:\Source\DomainManagement\DomainManagement\spiders\selector.json', 'r') as j:
            list_tag = json.loads(j.read())
    def GetContent(self, url):
        list_tag = self.list_tag
        for p in list_tag:
            req = requests.get('https://en.vneconomy.vn/')
            html = req.text
            soup = BeautifulSoup(html, 'html.parser')
            LstLink = soup.find('section' , class_="zone")
            list_of_company_a = LstLink.find_all("article")
            for s in list_of_company_a:
                try:
                    # title = s.find(p['title'], p['title_class']).text
                    title = s.find(p['title']).text
                    text = s.find(p['text']).text
                    time = s.find(p['time']).text
                    img = s.find("img")
                    link_image = img[p['img']]
                    img_head = requests.get(link_image)
                    with open("img_head.png", "wb") as im_h:
                        im_h.write(img_head.content)
                    if '' != title:
                        content.append({
                             title,
                             text,
                             #time,
                        })


                except:
                    pass

        return content
    def main(self):
        Idx = 0;
        len_list = len(list_child_url);
        while Idx < len_list:
            self.Extract_Url(list_child_url[Idx].url,list_child_url[Idx].deep)
            Idx = Idx + 1;
            len_list = len(list_child_url)


list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
object.main()
document = Document()
object.GetContent(url)
for obj in list_child_url:
    document.add_paragraph(obj.url)
    print(obj.url, obj.iscan, obj.deep, sep=' ')
print(len(list_child_url))
for value in content:
    document.add_paragraph(value)
    document.add_picture('img_head.png')
    print(value)
print(len(content))
# Vị trí lưu file
filePath = '/DomainManagement\DomainManagement\spiders/url.docx'
document.save("E:\Source" + filePath)

#up_load file lên server:
url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
#url_upload = resp_json['DataStoragePath']
response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "phancuoc_cntt_3i"}, files={
    "fileUpload": (
    'url.docx', open("E:\Source" + filePath, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
if response_upload.ok:
    print("Upload completed successfully!")
    print(response_upload.text)
else:
    print("Something went wrong!")

