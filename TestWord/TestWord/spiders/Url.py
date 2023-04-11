import json
from hashlib import new

import requests
from bs4 import BeautifulSoup
from docx import Document


def url_obj(url, param, param1):
    pass


class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep
    Idx = 0;
    chk = 0;
    def chk_link_exist(self, link):
        for obj in list_child_url:
            if (obj.url == link):
                 return 1;
        return 0;

    def Extract_Url(self,url,deep):
        try:
            if deep <= 1:
                req = requests.get(url)
                html = req.text
                soup = BeautifulSoup(html, 'html.parser')
                LstLink = soup.find('body')
                for s in LstLink.find_all('a'):
                    try:
                        link = s['href']
                        if 'https' not in link:
                            link = 'https://fptshop.com.vn' + link
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        elif 'https://' in link:
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                    except:
                        pass
                else:
                    pass
        except:
            print(1)
            pass
    with open(r'..\selector.json', 'r') as j:
       #C:\Users\pycha\PycharmProjects\Source\TestWord\TestWord\selector.json
            list_tag = json.loads(j.read())
    def GetContent(self, url):
        list_tag = self.list_tag
        for p in list_tag:

            req = requests.get(url)
            html = req.text
            soup = BeautifulSoup(html, 'html.parser')
            LstLink = soup.find('html')
            for s in LstLink.find_all('a'):
                try:

                    title = s[p['selector']]
                    ct = s[p['content']]

                    if '' != title:
                        content.append({
                           "selector": title,
                           "content": ct,
                        })

                except:
                    pass

            return content
    def main(self):
        Idx = 0;
        len_list = len(list_child_url);
        while Idx < len_list:
            self.Extract_Url(list_child_url[Idx].url,list_child_url[Idx].deep)
            Idx = Idx + 1;
            len_list = len(list_child_url)

list_child_url = []
content = []

url = "https://fptshop.com.vn/"
list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
#object.main()
#object.Extract_Url(url,1)
document = Document()
object.GetContent(url)
#for obj in list_child_url:
    #document.add_paragraph(obj.url)
    #print(obj.url, obj.iscan, obj.deep, sep=' ')
#print(len(list_child_url))
for title in content:
        document.add_paragraph(title)
        print(title)
print(len(content))

#document.add_heading("Time Start : " + dt_string, 2)
filePath  = 'Link_url.docx'
document.save("E:\Source\TestWord\\" + filePath)


