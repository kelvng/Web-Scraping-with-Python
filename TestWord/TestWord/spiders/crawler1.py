from cmath import log
from urllib import response
import pandas as pd
from scrapy import Spider
from scrapy.selector import Selector
import scrapy
import requests
from bs4 import BeautifulSoup
import urllib.request
import socket
from datetime import datetime
from docx import Document
import json
# from scrapy.settings import Settings
# from twisted.internet import reactor
from tqdm import tqdm

from .items import TestWordItem
from scrapy.crawler import CrawlerProcess
import ftplib


class CrawlerSpider(Spider):
    resp = requests.post("https://os.3i.com.vn/PythonCrawler/GetCrawlerData?spiderName=crawler")
    resp_json = resp.json()
    print(resp_json)
    print("Url")
    print(resp_json['Url'])

    name ="crawler"
    #name = resp_json['SpiderName']

    allowed_domains = ["en.vneconomy.vn"]
    # start_urls=["https://en.vneconomy.vn/"]
    start_urls = resp_json['Url']

    with open('..\data.json', 'r') as j:
        list_tag = json.loads(j.read())
    start_page = 1
    end_page = 100
    step_page = 50

    def parse(self, response):
        main_link = self.start_urls[0]
        list_tag = self.list_tag
        Locals = socket.gethostbyname(self.start_urls)
        now = datetime.now()
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        begin = "Start with website  " + main_link + "  Local IP  = " + Locals + "  Time Start : " + dt_string
        document = Document()
        document.add_heading("Start with website  " + main_link, 0)
        document.add_heading("Time Start : " + dt_string, 2)
        document.save("text1.docx")
        items = TestWordItem()
        items['begin'] = begin
        yield items
        questions = Selector(response).xpath(list_tag["list_group"].replace("'", '"'))
        print("a2")
        print(questions)
        for question in questions:
            tag_main = main_link + question.xpath(list_tag["link_group"].replace("'", '"')).extract_first()
            print("a3")
            print(tag_main)
            response = requests.get(tag_main)
            list_post_hl = Selector(response).xpath(list_tag["list_post_hl"].replace("'", '"'))
            print("a4")
            print(list_post_hl)
            for post_hl in list_post_hl:
                post = main_link + post_hl.xpath(list_tag["link_post_hl"].replace("'", '"')).extract_first()
                print("a5")
                print(list_post_hl)
                request = scrapy.Request(post, callback=self.parse_full_post)
                request.meta['post'] = post
                yield request
            start_page = self.start_page
            end_page = self.end_page
            step_page = self.step_page
            i = start_page
            while i < end_page:
                temp = tag_main + "?trang=" + str(i)
                response = requests.get(temp)
                list_post_page = Selector(response).xpath(list_tag["list_post"].replace("'", '"'))
                print("a6")
                print(list_tag["list_post"].replace("'", '"'))
                print(list_post_page)

                for post_page in list_post_page:
                    post = main_link + post_page.xpath(list_tag["link_post"].replace("'", '"')).extract_first()
                    print("a7")
                    print(post)
                    request = scrapy.Request(post, callback=self.parse_full_post)
                    yield request
                page = urllib.request.urlopen(temp)
                soup = BeautifulSoup(page, 'html.parser')
                new_feed = soup.find_all(list_tag["next_page"][0], class_=list_tag["next_page"][1])
                if new_feed == []:
                    i = 2000
                i += step_page

    def parse_full_post(self, response):
        list_tag = self.list_tag
        Full_Post = Selector(response)
        print(Full_Post)
        url_post = response.xpath(list_tag["url"].replace("'", '"')).extract_first()
        date = response.css(list_tag["date"].replace("'", '"')).get()
        title = response.css(list_tag["title"].replace("'", '"')).get()
        author = Selector(response).xpath(list_tag["author"].replace("'", '"')).extract_first()
        head = response.css(list_tag["head"].replace("'", '"')).get()

        document = Document("text1.docx")
        document.add_paragraph("__________________________________________________________________________")
        a = document.add_paragraph()
        b = a.add_run(url_post)
        b.italic = True
        b.underline = True
        b.font.name = 'Tahoma'
        # b.color.rgb = RGBColor(0, 0, 255)
        ########################
        a = document.add_paragraph()
        b = a.add_run(date.replace("GMT+7", ""))
        b.font.name = 'Tahoma'
        # url.font.size =12000
        a = document.add_paragraph()
        b = a.add_run(author)
        b.italic = True
        a = document.add_heading(title, 1)
        a.alignment = 1
        a = document.add_paragraph()
        b = a.add_run(head)
        b.both = True
        #image
        link_img_head = Selector(response).xpath(list_tag["image"]).extract_first()
        img_head = requests.get(link_img_head)
        with open("img_head.png", "wb") as im_h:
            im_h.write(img_head.content)
        a = document.add_picture('img_head.png')
        a.alignment = 1
        #content
        #contents = select(response)
        contents = Selector(response).xpath(list_tag["contents"].replace("'", '"')).get()
        soup = BeautifulSoup(contents, 'html.parser')
        i_p = 0
        i_h3 = 0
        i_f = 0
        for tag in soup.find_all(["p", "h3", "figure"]):
            if tag.name == "p":
                p = soup.find_all("p")[i_p].string
                document.add_paragraph(p)
                i_p += 1
            if tag.name == "h3":
                h3 = soup.find_all("h3")[i_h3].string
                a = document.add_paragraph()
                b = a.add_run(h3).both = True
                b.alignment = 1
                i_h3 += 1
            if tag.name == "figure":
                link_img = soup.find_all("img")[i_f]['src']
                img = requests.get(link_img)
                with open("img.png", "wb") as im:
                    im.write(img.content)
                a = document.add_picture('img.png')
                a.alignment = 1
                print(soup.find_all("img")[i_f]['alt'])
                i_f += 1
        document.save("text1.docx")
        item = TestWordItem()
        item['url_post'] = url_post
        item['title'] = title
        item['date'] = date
        item['author'] = author
        item['head'] = head
        yield item


process = CrawlerProcess()
process.crawl(CrawlerSpider)
process.start()

# ftp = ftplib.FTP("117.6.131.222")
# ftp.login("crawler", "kiet1234")
# localfile='text.docx'
# remotefile='text.docx'
# with open(localfile, "rb") as file:
#     ftp.storbinary('STOR %s' % remotefile, file)

url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "dunglequoc_cntt_3i"}, files={
    "fileUpload": (
    'text1.docx', open('text1.docx', 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
if response_upload.ok:
    print("Upload completed successfully!")
    print(response_upload.text)
else:
    print("Something went wrong!")