# coding=utf8
# -*- coding: utf-8 -*-
import ast
import json
import os
from time import sleep
import docx
import scrapy
import websockets
import asyncio
import sys
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime

from docx.enum.dml import MSO_THEME_COLOR_INDEX
time = datetime.now()
start_time = time.strftime("%H:%M:%S")
print("Current Time =", start_time)

list_child_url = []
url_scan = []
# mảng chứa lisr url và content text abc
content = []
scancontenturl = []
Lstkeyword = []
botcode = []

list_child_url1 = []
url_scan1 = []
# mảng chứa lisr url và content text abc
content1 = []
scancontenturl1 = []
Lstkeyword1 = []
botcode1 = []

list_child_url2 = []
url_scan2 = []
# mảng chứa lisr url và content text abc
content2 = []
scancontenturl2 = []
Lstkeyword2 = []
botcode2 = []

URL = 'ws://127.0.0.1:9091'


def outer_func():
    class url_obj:

        def __init__(self, url, iscan, deep, selector, keyword, deepscan, linkadd):
            self.url = url
            self.iscan = iscan
            self.deep = deep
            self.selector = selector
            self.keyword = keyword
            self.deepscan = deepscan
            self.linkadd = linkadd

        Idx = 0
        chk = 0

        # Hàm kiểm tra link đã tồn tại trong list chưa
        def chk_link_exist(self, link):
            for obj in list_child_url:
                if (obj.url == link):
                    return 1
            return 0

        # Hàm quét và lấy các link con từ link mẹ
        def Extract_Url(self, url, deep, deepscan, linkadd):
            try:
                # Độ sâu của link mẹ là 1, link con = link mẹ +1
                if deep <= deepscan:
                    self.deep = deep
                    req = requests.get(url)
                    html = req.text
                    soup = BeautifulSoup(html, 'html.parser')
                    LstLink = soup.find('body')
                    for s in LstLink.find_all('a'):
                        try:
                            link = s['href']
                            if 'http' not in link:
                                link = linkadd + link

                                if self.chk_link_exist(link) == 0:
                                    url = link
                                    iscan = 0
                                    deep = self.deep + 1
                                    list_child_url.append(
                                        url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                self.linkadd))
                            else:
                                if linkadd in link:
                                    if self.chk_link_exist(link) == 0:
                                        url = link
                                        iscan = 0
                                        deep = self.deep + 1
                                        list_child_url.append(
                                            url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                    self.linkadd))
                                else:
                                    print('Strange link!')
                                    pass
                        except:
                            pass
            except:
                print('link fail')
                pass

        def add_hyperlink(self, paragraph, text, url):
            # This gets access to the document.xml.rels file and gets a new relation id value
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
            # Create a w:r element and a new w:rPr element
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            # Create a new Run object and add the hyperlink into it
            r = paragraph.add_run()
            r._r.append(hyperlink)
            # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
            # Delete this if using a template that has the hyperlink style in it
            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True
            return hyperlink

        def GetContent(self, url, selector, keyword):
            try:
                req = requests.get(url)

                html = req.text
                # html = html.decode('cp1251')
                soup = BeautifulSoup(html, 'lxml')
                header_local = soup.find(selector["header"], class_=selector["headerclass"])
                header = header_local.text
                # print(header)
                Content_local = soup.find(selector["content"], class_=selector["contentclass"])
                # content.append(keyword)
                for img in soup.findAll(True, selector['imageclass']):
                    # for img in Content_local.find_all(selector['image']):
                    image = img.find('img')['src']
                Content = Content_local.text
                for x in keyword:
                    if x in header:
                        content.append('Link post: ' + url)
                        scancontenturl.append(url)
                        Url = url + " status: " + "True"
                        url_scan.append(url)
                        content.append(header)
                        print(Url)
                        content.append(Content)
                        # content.append(image)
                    else:
                        Url = url + " status: " + "False"
                        url_scan.append(Url)
                return content
            except:
                pass

        def save_file(self, document):
            filename = botcode[0] + '.docx'
            file_path = filename
            document.save(file_path)
            # r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\save\url1.docx')
            url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
            # url_upload = resp_json['DataStoragePath']

            response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "huynv_cntt_3i"},
                                            files={
                                                "fileUpload": (
                                                    filename,
                                                    open(
                                                        file_path,
                                                        'rb'),
                                                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
            if response_upload.ok:
                print("Upload completed successfully!")
                print(response_upload.text)
                sleep(4)
                os.remove(file_path)
            else:
                print("Something went wrong!")

        async def main(self):
            Idx = 0
            len_list = len(list_child_url)
            await asyncio.sleep(10)
            while Idx < len_list:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    try:
                        self.Extract_Url(list_child_url[Idx].url, list_child_url[Idx].deep,
                                         list_child_url[Idx].deepscan, list_child_url[Idx].linkadd)
                    except:
                        pass
                    try:
                        self.GetContent(list_child_url[Idx].url, list_child_url[Idx].selector,
                                        list_child_url[Idx].keyword)
                    except:
                        pass
                    await ws.send(str(datetime.now()) + ": " + list_child_url[Idx].url)
                    Idx = Idx + 1
                    len_list = len(list_child_url)
                    print(len_list)
                    print(Idx)
                    pass

            print(1)
            time = datetime.now()
            end_time = time.strftime("%H:%M:%S")
            urlscan_json = json.dumps(url_scan)
            # data["UrlScanJson"] = urlscan_json
            # data["endtime"] = end_time

            data = {
                'SessionCode': botcode[0] + end_time,
                'StartTime': start_time,
                'EndTime': end_time,
                'UrlScanJson': urlscan_json,
                'FileDownloadJson': 'url1.docx',
                'NumOfFile': 1,
                'FileResultData': '',
                'NumPasscap': '',
                'UserIdRunning': '001',
                'Ip': '1',
                'Status': 'active',
                'BotCode': botcode[0],
                'TimeScan': '',
                'CreatedBy': 'admin',
            }
            url_upload = "https://dieuhanh.vatco.vn/PythonCrawler/InsertCrawlerRunningLog"
            resp = requests.post(url_upload, data=data)
            if resp.ok:
                print("Upload completed successfully!")
                print(resp.text)
            else:
                print("Something went wrong!")
            document = Document()
            document.add_paragraph('Search With Keyword:' + str(Lstkeyword))
            print('Done. Start save text!')
            for value in content:
                p = document.add_paragraph(value)
                for url in scancontenturl:
                    if url in value:
                        self.add_hyperlink(p, 'Link!', url)
                        # self.delete_paragraph(p)
            self.save_file(document)
            content.clear()
            list_child_url.clear()
            url_scan.clear()
            scancontenturl.clear()
            Lstkeyword.clear()
            botcode.clear()
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Write file and Finish!')
    class url_obj1:

        def __init__(self, url, iscan, deep, selector, keyword, deepscan, linkadd):
            self.url = url
            self.iscan = iscan
            self.deep = deep
            self.selector = selector
            self.keyword = keyword
            self.deepscan = deepscan
            self.linkadd = linkadd

        Idx = 0
        chk = 0

        # Hàm kiểm tra link đã tồn tại trong list chưa
        def chk_link_exist(self, link):
            for obj in list_child_url1:
                if (obj.url == link):
                    return 1
            return 0

        # Hàm quét và lấy các link con từ link mẹ
        def Extract_Url(self, url, deep, deepscan, linkadd):
            try:
                # Độ sâu của link mẹ là 1, link con = link mẹ +1
                if deep <= deepscan:
                    self.deep = deep
                    req = requests.get(url)
                    html = req.text
                    soup = BeautifulSoup(html, 'html.parser')
                    LstLink = soup.find('body')
                    for s in LstLink.find_all('a'):
                        try:
                            link = s['href']
                            if 'http' not in link:
                                link = linkadd + link

                                if self.chk_link_exist(link) == 0:
                                    url = link
                                    iscan = 0
                                    deep = self.deep + 1
                                    list_child_url1.append(
                                        url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                self.linkadd))
                            else:
                                if linkadd in link:
                                    if self.chk_link_exist(link) == 0:
                                        url = link
                                        iscan = 0
                                        deep = self.deep + 1
                                        list_child_url1.append(
                                            url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                    self.linkadd))
                                else:
                                    print('Strange link!')
                                    pass
                        except:
                            pass
            except:
                print('link fail')
                pass

        def add_hyperlink(self, paragraph, text, url):
            # This gets access to the document.xml.rels file and gets a new relation id value
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
            # Create a w:r element and a new w:rPr element
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            # Create a new Run object and add the hyperlink into it
            r = paragraph.add_run()
            r._r.append(hyperlink)
            # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
            # Delete this if using a template that has the hyperlink style in it
            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True
            return hyperlink

        def GetContent(self, url, selector, keyword):
            try:
                req = requests.get(url)

                html = req.text
                # html = html.decode('cp1251')
                soup = BeautifulSoup(html, 'lxml')
                header_local = soup.find(selector["header"], class_=selector["headerclass"])
                header = header_local.text
                # print(header)
                Content_local = soup.find(selector["content"], class_=selector["contentclass"])
                # content.append(keyword)
                for img in soup.findAll(True, selector['imageclass']):
                    # for img in Content_local.find_all(selector['image']):
                    image = img.find('img')['src']
                Content = Content_local.text
                for x in keyword:
                    if x in header:
                        content1.append('Link post: ' + url)
                        scancontenturl1.append(url)
                        Url = url + " status: " + "True"
                        url_scan1.append(url)
                        content1.append(header)
                        print(Url)
                        content1.append(Content)
                        # content.append(image)
                    else:
                        Url = url + " status: " + "False"
                        url_scan1.append(Url)
                return content1
            except:
                pass

        def save_file(self, document):
            filename = botcode1[0] + '.docx'
            file_path = filename
            document.save(file_path)
            # r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\save\url1.docx')
            url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
            # url_upload = resp_json['DataStoragePath']

            response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "huynv_cntt_3i"},
                                            files={
                                                "fileUpload": (
                                                    filename,
                                                    open(
                                                        file_path,
                                                        'rb'),
                                                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
            if response_upload.ok:
                print("Upload completed successfully!")
                print(response_upload.text)
                sleep(4)
                os.remove(file_path)
            else:
                print("Something went wrong!")

        async def main(self):
            Idx = 0
            len_list = len(list_child_url1)
            await asyncio.sleep(10)
            while Idx < len_list:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    try:
                        self.Extract_Url(list_child_url1[Idx].url, list_child_url1[Idx].deep,
                                         list_child_url1[Idx].deepscan, list_child_url1[Idx].linkadd)
                    except:
                        pass
                    try:
                        self.GetContent(list_child_url1[Idx].url, list_child_url1[Idx].selector,
                                        list_child_url1[Idx].keyword)
                    except:
                        pass
                    await ws.send(str(datetime.now()) + ": " + list_child_url1[Idx].url)
                    Idx = Idx + 1
                    len_list = len(list_child_url1)
                    print(len_list)
                    print(Idx)
                    pass

            print(1)
            time = datetime.now()
            end_time = time.strftime("%H:%M:%S")
            urlscan_json = json.dumps(url_scan1)
            # data["UrlScanJson"] = urlscan_json
            # data["endtime"] = end_time

            data = {
                'SessionCode': botcode1[0] + end_time,
                'StartTime': start_time,
                'EndTime': end_time,
                'UrlScanJson': urlscan_json,
                'FileDownloadJson': 'url1.docx',
                'NumOfFile': 1,
                'FileResultData': '',
                'NumPasscap': '',
                'UserIdRunning': '001',
                'Ip': '1',
                'Status': 'active',
                'BotCode': botcode1[0],
                'TimeScan': '',
                'CreatedBy': 'admin',
            }
            url_upload = "https://dieuhanh.vatco.vn/PythonCrawler/InsertCrawlerRunningLog"
            resp = requests.post(url_upload, data=data)
            if resp.ok:
                print("Upload completed successfully!")
                print(resp.text)
            else:
                print("Something went wrong!")
            document = Document()
            document.add_paragraph('Search With Keyword:' + str(Lstkeyword1))
            print('Done. Start save text!')
            for value in content:
                p = document.add_paragraph(value)
                for url in scancontenturl1:
                    if url in value:
                        self.add_hyperlink(p, 'Link!', url)
                        # self.delete_paragraph(p)
            self.save_file(document)
            content.clear()
            list_child_url1.clear()
            url_scan1.clear()
            scancontenturl1.clear()
            Lstkeyword1.clear()
            botcode1.clear()
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Write file and Finish!')
    class url_obj2:

        def __init__(self, url, iscan, deep, selector, keyword, deepscan, linkadd):
            self.url = url
            self.iscan = iscan
            self.deep = deep
            self.selector = selector
            self.keyword = keyword
            self.deepscan = deepscan
            self.linkadd = linkadd

        Idx = 0
        chk = 0

        # Hàm kiểm tra link đã tồn tại trong list chưa
        def chk_link_exist(self, link):
            for obj in list_child_url:
                if (obj.url == link):
                    return 1
            return 0

        # Hàm quét và lấy các link con từ link mẹ
        def Extract_Url(self, url, deep, deepscan, linkadd):
            try:
                # Độ sâu của link mẹ là 1, link con = link mẹ +1
                if deep <= deepscan:
                    self.deep = deep
                    req = requests.get(url)
                    html = req.text
                    soup = BeautifulSoup(html, 'html.parser')
                    LstLink = soup.find('body')
                    for s in LstLink.find_all('a'):
                        try:
                            link = s['href']
                            if 'http' not in link:
                                link = linkadd + link

                                if self.chk_link_exist(link) == 0:
                                    url = link
                                    iscan = 0
                                    deep = self.deep + 1
                                    list_child_url.append(
                                        url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                self.linkadd))
                            else:
                                if linkadd in link:
                                    if self.chk_link_exist(link) == 0:
                                        url = link
                                        iscan = 0
                                        deep = self.deep + 1
                                        list_child_url.append(
                                            url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                    self.linkadd))
                                else:
                                    print('Strange link!')
                                    pass
                        except:
                            pass
            except:
                print('link fail')
                pass

        def add_hyperlink(self, paragraph, text, url):
            # This gets access to the document.xml.rels file and gets a new relation id value
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
            # Create a w:r element and a new w:rPr element
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            # Create a new Run object and add the hyperlink into it
            r = paragraph.add_run()
            r._r.append(hyperlink)
            # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
            # Delete this if using a template that has the hyperlink style in it
            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True
            return hyperlink

        def GetContent(self, url, selector, keyword):
            try:
                req = requests.get(url)

                html = req.text
                # html = html.decode('cp1251')
                soup = BeautifulSoup(html, 'lxml')
                header_local = soup.find(selector["header"], class_=selector["headerclass"])
                header = header_local.text
                # print(header)
                Content_local = soup.find(selector["content"], class_=selector["contentclass"])
                # content.append(keyword)
                for img in soup.findAll(True, selector['imageclass']):
                    # for img in Content_local.find_all(selector['image']):
                    image = img.find('img')['src']
                Content = Content_local.text
                for x in keyword:
                    if x in header:
                        content.append('Link post: ' + url)
                        scancontenturl.append(url)
                        Url = url + " status: " + "True"
                        url_scan.append(url)
                        content.append(header)
                        print(Url)
                        content.append(Content)
                        # content.append(image)
                    else:
                        Url = url + " status: " + "False"
                        url_scan.append(Url)
                return content
            except:
                pass

        def save_file(self, document):
            filename = botcode[0] + '.docx'
            file_path = filename
            document.save(file_path)
            # r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\save\url1.docx')
            url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
            # url_upload = resp_json['DataStoragePath']

            response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "huynv_cntt_3i"},
                                            files={
                                                "fileUpload": (
                                                    filename,
                                                    open(
                                                        file_path,
                                                        'rb'),
                                                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
            if response_upload.ok:
                print("Upload completed successfully!")
                print(response_upload.text)
                sleep(4)
                os.remove(file_path)
            else:
                print("Something went wrong!")

        async def main(self):
            Idx = 0
            len_list = len(list_child_url)
            await asyncio.sleep(10)
            while Idx < len_list:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    try:
                        self.Extract_Url(list_child_url[Idx].url, list_child_url[Idx].deep,
                                         list_child_url[Idx].deepscan, list_child_url[Idx].linkadd)
                    except:
                        pass
                    try:
                        self.GetContent(list_child_url[Idx].url, list_child_url[Idx].selector,
                                        list_child_url[Idx].keyword)
                    except:
                        pass
                    await ws.send(str(datetime.now()) + ": " + list_child_url[Idx].url)
                    Idx = Idx + 1
                    len_list = len(list_child_url)
                    print(len_list)
                    print(Idx)
                    pass

            print(1)
            time = datetime.now()
            end_time = time.strftime("%H:%M:%S")
            urlscan_json = json.dumps(url_scan)
            # data["UrlScanJson"] = urlscan_json
            # data["endtime"] = end_time

            data = {
                'SessionCode': botcode[0] + end_time,
                'StartTime': start_time,
                'EndTime': end_time,
                'UrlScanJson': urlscan_json,
                'FileDownloadJson': 'url1.docx',
                'NumOfFile': 1,
                'FileResultData': '',
                'NumPasscap': '',
                'UserIdRunning': '001',
                'Ip': '1',
                'Status': 'active',
                'BotCode': botcode[0],
                'TimeScan': '',
                'CreatedBy': 'admin',
            }
            url_upload = "https://dieuhanh.vatco.vn/PythonCrawler/InsertCrawlerRunningLog"
            resp = requests.post(url_upload, data=data)
            if resp.ok:
                print("Upload completed successfully!")
                print(resp.text)
            else:
                print("Something went wrong!")
            document = Document()
            document.add_paragraph('Search With Keyword:' + str(Lstkeyword))
            print('Done. Start save text!')
            for value in content:
                p = document.add_paragraph(value)
                for url in scancontenturl:
                    if url in value:
                        self.add_hyperlink(p, 'Link!', url)
                        # self.delete_paragraph(p)
            self.save_file(document)
            content.clear()
            list_child_url.clear()
            url_scan.clear()
            scancontenturl.clear()
            Lstkeyword.clear()
            botcode.clear()
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Write file and Finish!')

    async def initspider(url, selector, keyword, deepscan, linkadd):
        list_child_url.append(url_obj(url, 0, 1, selector, keyword, deepscan, linkadd))
        object = url_obj(url, 0, 1, selector, keyword, deepscan, linkadd)
        await asyncio.sleep(5)
        # data['keyword'] = keyword
        for x in keyword:
            Lstkeyword.append(x)
        asyncio.create_task(object.main())

    async def listen():
        ws_connect = websockets.connect('ws://127.0.0.1:9091', ping_interval=None)
        my_task = None
        async with ws_connect as wb:
            await wb.send('Spider running!')
            while True:
                param = await wb.recv()
                if "Url" in param:
                    data = json.loads(param)
                    print(data)
                    deepscan = data['DeepScan']
                    param1 = data['Url']
                    selector = data['ConfigSelectorJson']
                    selector = json.loads(selector)
                    checkwordlist = data['ListKeyWord']
                    checkwordlist = ast.literal_eval(checkwordlist)
                    print(checkwordlist)
                    BotCode = data['BotCode']
                    botcode.append(BotCode)
                    addurl = data['Url']
                    my_task = asyncio.create_task(initspider(param1, selector, checkwordlist, deepscan, addurl))
                if 'Stop domain' in param:
                    # flagstop = True
                    # sys.exit()
                    if my_task:
                        print('Stop task')
                        my_task.cancel()
                        break
                    # content.clear()
                    # list_child_url.clear()
                    # url_scan.clear()
                    # scancontenturl.clear()
                    # Lstkeyword.clear()
                    # botcode.clear()
                print("hearing")
            print("Stop async with ws_connect")
        print('Stop listen function')

    async def main():
        task_1 = asyncio.create_task(listen())
        await asyncio.sleep(0.250)
        await task_1

    # async def forever():
    #     while True:
    #         await main()
    loop = asyncio.get_event_loop()

    def startTask():
        try:
            loop.run_until_complete(main())
        except asyncio.CancelledError:
            print("some error but still fine")
        finally:
            print("live again")
            startTask()

    startTask()


class Myspider(scrapy.Spider):
    name = 'CoreWebSpider'
    some_attribute = "Yes|No"
    print("hello world !")
    outer_func()




