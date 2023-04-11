# coding=utf8
# -*- coding: utf-8 -*-
import ast
import json
import os
import uuid
from concurrent.futures.thread import ThreadPoolExecutor
from time import sleep
import docx
import scrapy
import websockets
import asyncio
import sys
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime

from docx.enum.dml import MSO_THEME_COLOR_INDEX

CHECK = True
list_child_url = []
url_scan = []
# mảng chứa lisr url và content text abc
content = []
list_str = []
scancontenturl = []
Lstkeyword = []
botcode = []
List_search = []
URL = 'ws://127.0.0.1:9091'

def outer_func():
    class url_obj:
        def __init__(self, url, iscan, deep, selector, keyword, deepscan, linkadd, StartTime):
            self.url = url
            self.iscan = iscan
            self.deep = deep
            self.selector = selector
            self.keyword = keyword
            self.deepscan = deepscan
            self.linkadd = linkadd
            self.StartTime = StartTime

        Idx = 0
        chk = 0

        # Hàm kiểm tra link đã tồn tại trong list chưa
        def chk_link_exist(self, link):
            for obj in list_child_url:
                if (obj.url == link):
                    return 1
            return 0

        # Hàm quét và lấy các link con từ link mẹ
        def Extract_Url(self, url, deep, deepscan, linkadd):
            try:
                # Độ sâu của link mẹ là 1, link con = link mẹ +1
                if deep <= deepscan:
                    self.deep = deep
                    req = requests.get(url)
                    html = req.text
                    print(url)
                    soup = BeautifulSoup(html, 'html.parser')
                    LstLink = soup.find('body')
                    for s in LstLink.find_all('a'):
                        try:
                            link = s['href']
                            if 'http' not in link:
                                link = linkadd + link

                                if self.chk_link_exist(link) == 0:
                                    url = link
                                    iscan = 0
                                    deep = self.deep + 1
                                    list_child_url.append(
                                        url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                self.linkadd, self.StartTime))
                            else:
                                if linkadd in link:
                                    if self.chk_link_exist(link) == 0:
                                        url = link
                                        iscan = 0
                                        deep = self.deep + 1
                                        list_child_url.append(
                                            url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                    self.linkadd, self.StartTime))
                                else:
                                    pass
                        except:
                            pass
            except:
                print('link fail')
                pass

        def add_hyperlink(self, paragraph, text, url):

            # This gets access to the document.xml.rels file and gets a new relation id value
            part = paragraph.part
            # hyperlink
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
            # Create a w:r element and a new w:rPr element
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            # Create a new Run object and add the hyperlink into it
            r = paragraph.add_run()
            r._r.append(hyperlink)
            # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
            # Delete this if using a template that has the hyperlink style in it
            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True
            print("Hyperlink type" + str(type(hyperlink)))

            return hyperlink

        def GetContent(self, url, selector, keyword):
            try:
                List_search.clear()
                list_str.clear()
                req = requests.get(url)
                html = req.text
                # html = html.decode('cp1251')
                soup = BeautifulSoup(html, 'lxml')
                header_local = soup.find(selector["header"], class_=selector["headerclass"])
                header = header_local.text
                # print(header)
                Content_local = soup.find(selector["content"], class_=selector["contentclass"])
                Content = Content_local.text

                List_count = {
                    'Key': '',
                    'Times': 1
                }
                linkpostdata = ''
                contentdata = ''
                headerdata = ''
                if any(x in Content for x in keyword):

                    for i in keyword:
                        List_count['Key'] = i
                        List_count['Times'] = Content.count(i) + header.count(i)
                        if List_count['Times'] == 0:
                            print("not have time----")
                            continue
                        else:
                            print("have content")

                            op = json.dumps(List_count)
                            print(op)
                            List_search.append(op)

                    for j in List_search:
                        data = json.loads(j)
                        print(data)
                        list_str.append(data)

                    scancontenturl.append(url)
                    print(list_str)
                    # list_str2 = list_str.replace("'",'"')
                    list_str1 = json.dumps(list_str)
                    print(list_str1)
                    Url = url
                    scan_json = {
                        'url': Url,
                        'keyword': list_str1,
                        'Status': True
                    }
                    print(scan_json)
                    linkpostdata = 'Link post: ' + url
                    headerdata = header
                    contentdata = Content
                    url_scan.append(scan_json)

                if linkpostdata == '' and contentdata == '' and headerdata == '':
                    Url = url
                    scan_json = {
                        'url': Url,
                        'keyword': list_str,
                        'Status': False
                    }
                    url_scan.append(scan_json)
                if linkpostdata != '' and contentdata != '' and headerdata != '' and str(list_str) != '':
                    content.append(linkpostdata)
                    content.append("Keywords search:" + str(List_search))
                    content.append(headerdata)
                    content.append(contentdata)
                return content
            except:
                pass

        def upload_data(self, document):
            timebegin = self.StartTime
            print(timebegin)
            start_time1 = timebegin.strftime("%H:%M:%S")

            myUUID = uuid.uuid4()
            mystr = str(myUUID)
            print(botcode)
            filename = botcode[0] + '-' +mystr + '.docx'
            file_path = filename
            document.save(file_path)

            filesize = os.path.getsize(file_path)
            # r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\save\url1.docx')
            url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
            # url_upload = resp_json['DataStoragePath']

            response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "huynv_cntt_3i"},
                                            files={
                                                "fileUpload": (
                                                    filename,
                                                    open(
                                                        file_path,
                                                        'rb'),
                                                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
            if response_upload.ok:
                print("Upload completed successfully!")
                print("Tải file lên DB!")
                print(response_upload.text)
                sleep(4)
                os.remove(file_path)
                statusupload = response_upload.text
                # {"ID":22577,"Title":"Thêm mới tệp tin thành công","Error":false,"Object":{"Type":"SERVER","Url":"https://dieuhanh.vatco.vn//uploads/repository/CRAWLING/vneconomy.docx","CloudId":""},"Code":null}
                datasend = json.loads(statusupload)
                File_Result = datasend['Object']["Url"]
            else:
                print("Something went wrong!")
                File_Result = ''

            end_time = datetime.now()
            print(end_time)
            end_time1 = end_time.strftime("%H:%M:%S")
            urlscan_json = json.dumps(url_scan)
            print(urlscan_json)
            # data["UrlScanJson"] = urlscan_json
            # data["endtime"] = end_time
            print(Lstkeyword)
            data = {
                'SessionCode': botcode[0] + '_' + end_time1,
                'StartTime': start_time1,
                'EndTime': end_time1,
                'UrlScanJson': urlscan_json,
                'FileDownloadJson': '',
                'NumOfFile': 1,
                'FileResultData': File_Result,
                'NumPasscap': '',
                'UserIdRunning': '001',
                'Ip': '1',
                'Status': 'active',
                'BotCode': botcode[0],
                'TimeScan': int((end_time - timebegin).total_seconds()),
                'CreatedBy': 'Huycntt_3i',
                'KeyWord': str(Lstkeyword),
                'FileSizeDownload': int(filesize),
            }
            # [FILE_SIZE_DOWNLOAD
            print(data)
            url_upload = "http://localhost:6023/PythonCrawler/InsertCrawlerRunningLog"
            # url_upload = "https://dieuhanh.vatco.vn/PythonCrawler/InsertCrawlerRunningLog"
            resp = requests.post(url_upload, data=data)
            if resp.ok:
                print("Upload completed successfully!")
                print("data report crawl!")
                print(resp.text)
            else:
                print("Something went wrong!")

        async def main(self):

            Idx = 0
            len_list = len(list_child_url)
            await asyncio.sleep(10)
            while Idx < len_list:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    try:
                        self.Extract_Url(list_child_url[Idx].url, list_child_url[Idx].deep,
                                         list_child_url[Idx].deepscan, list_child_url[Idx].linkadd)
                    except:
                        pass
                    try:
                        self.GetContent(list_child_url[Idx].url, list_child_url[Idx].selector,
                                        list_child_url[Idx].keyword)
                        # print(list_child_url[Idx].url)
                        print("start search content")
                    except:
                        pass
                    await ws.send(str(datetime.now()) + ": " + list_child_url[Idx].url)
                    Idx = Idx + 1
                    len_list = len(list_child_url)
                    pass
            document = Document()
            document.add_paragraph('Search With Keyword:' + str(Lstkeyword))
            print('Done. Start save text!')
            for value in content:
                p = document.add_paragraph(value)
                for url in scancontenturl:

                    if url in value:
                        self.add_hyperlink(p, 'Link!', url)

                        # self.delete_paragraph(p)
            self.upload_data(document)
            document.save(r'F:\PycharmProjects\Source\DomainManagement\DomainManagement\spiders\save/' + 'test.docx')
            # print(1)

            content.clear()
            list_child_url.clear()
            url_scan.clear()
            scancontenturl.clear()
            Lstkeyword.clear()
            botcode.clear()
            list_str.clear()

            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Write file and Finish!')

    async def initspider1(url, selector, keyword, deepscan, linkadd):
        start_time = datetime.now()
        start_time1 = start_time.strftime("%H:%M:%S")
        print("Current Time =", start_time1)
        for x in keyword:
            Lstkeyword.append(x)
        object = url_obj(url, 0, 1, selector, keyword, deepscan, linkadd, start_time)
        list_child_url.append(object)
        task = asyncio.create_task(object.main())
        await asyncio.gather(task)

        print('jobdone')

    async def takeparam(param):
        dataparam = json.loads(param)
        thread = []
        for data in dataparam:

            deepscan = data['DeepScan']
            param1 = data['Url']
            selector = data['ConfigSelectorJson']
            selector = json.loads(selector)
            print(selector)
            checkwordlist = data['ListKeyWord']
            checkwordlist = ast.literal_eval(checkwordlist)
            print(checkwordlist)
            start_time = datetime.now()
            start_time1 = start_time.strftime("%H:%M:%S")
            print("Current Time =", start_time1)
            for x in checkwordlist:
                Lstkeyword.append(x)

            BotCode = data['BotCode']
            botcode.append(BotCode)
            addurl = data['Url']
            object = url_obj(param1, 0, 1, selector, checkwordlist, deepscan, addurl, start_time)
            list_child_url.append(object)
            task = object.main()
            thread.append(task)
        print(thread)
        with ThreadPoolExecutor(max_workers=10) as executor:
           for task in thread:
                task1 = executor.map(task,Iterable = True,timeout= 1,chunksize=1024)

                # await asyncio.gather(task1)
                # future  = executor.submit(task)
                # results = future.result()
                # print(results)
                task2 = asyncio.create_task(task1)
                #await asyncio.gather(task2)








    async def listen():

        print("oke1")
        ws_connect = websockets.connect('ws://127.0.0.1:9091', ping_interval=None)
        print("oke2")
        my_task = None
        task_param = None
        async with ws_connect as wb:
            await wb.send('Spider running!')
            while True:
                param = await wb.recv()
                print(param)
                if "Url" in param:
                    my_task = asyncio.create_task(takeparam(param))


                if 'Stop domain' in param:
                    # flagstop = True
                    # sys.exit()
                    if my_task:
                        print('Stop task')
                        my_task.cancel()
                        break
                    if task_param:
                        print('Stop task')
                        my_task.cancel()
                        break
            print("Stop async with ws_connect")
        print('Stop listen function')

    async def main():

        task_1 = asyncio.create_task(listen())
        await asyncio.sleep(0.250)
        await task_1

    # async def forever():
    #     while True:
    #         await main()
    loop = asyncio.get_event_loop()

    def startTask():
        try:
            loop.run_until_complete(main())
        except asyncio.CancelledError:
            print("some error but still fine")
        finally:
            print("live again")
            startTask()

    startTask()


class Myspider(scrapy.Spider):
    name = 'CoreWebSpider'
    some_attribute = "Yes|No"
    print("hello world !")
    outer_func()




