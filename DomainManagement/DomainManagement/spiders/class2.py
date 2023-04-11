import ast
import json
import os
import uuid
from time import sleep
import docx
import scrapy
import websockets
import asyncio
import sys
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from tqdm import tqdm
from os.path import exists

# mảng chứa lisr url và content text abc
# CHECK = True
# list_child_url = []
# url_scan = []
# content = []
# list_str = []
# scancontenturl = []
# Lstkeyword = []
# botcode = []
# List_search = []
# sesioncode = []
# Textcontent = []
URL = 'ws://127.0.0.1:9094'


class url_obj:
    # def __init__(self, url, iscan, deep, selector, keyword, deepscan, linkadd, StartTime, SessionCode, list_child_url,
    #              url_scan, content, list_str, scancontenturl, Lstkeyword, botcode, List_search, sesioncode,
    #              Textcontent):
    def __init__(self, url, iscan, deep, selector, keyword, deepscan, linkadd, StartTime, SessionCode, RobotCode):
        self.url = url
        self.iscan = iscan
        self.deep = deep
        self.selector = selector
        self.keyword = keyword
        self.deepscan = deepscan
        self.linkadd = linkadd
        self.StartTime = StartTime
        self.RobotCode = RobotCode
        self.SessionCode = SessionCode
        self.content=[]
        self.list_child_url = []
        self.url_scan = []
        self.list_str = []
        self.scancontenturl = []
        self.Lstkeyword = []
        self.botcode = []
        self.List_search = []
        self.sesioncode = []
        self.Textcontent = []

    Idx = 0
    chk = 0

    # Hàm kiểm tra link đã tồn tại trong list chưa
    def chk_link_exist(self, link):
        for obj in self.list_child_url:
            if obj.url == link:
                return 1
        return 0

    # Hàm quét và lấy các link con từ link mẹ
    def Extract_Url(self, url, deep, deepscan, linkadd):
        try:
            # Độ sâu của link mẹ là 1, link con = link mẹ +1
            if deep <= deepscan:
                self.deep = deep
                req = requests.get(url)
                html = req.text
                print(url)
                soup = BeautifulSoup(html, 'html.parser')
                LstLink = soup.find('body')
                for s in LstLink.find_all('a'):
                    try:
                        link = s['href']
                        if 'http' not in link:
                            link = linkadd + link

                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                self.list_child_url.append(
                                    url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                            self.linkadd, self.StartTime, self.SessionCode, self.RobotCode))
                                # list_child_url,url_scan,content,list_str,scancontenturl,Lstkeyword,botcode,
                                # List_search,sesioncode,Textcontent
                        else:
                            if linkadd in link:
                                if self.chk_link_exist(link) == 0:
                                    url = link
                                    iscan = 0
                                    deep = self.deep + 1
                                    self.list_child_url.append(
                                        url_obj(url, iscan, deep, self.selector, self.keyword, self.deepscan,
                                                self.linkadd, self.StartTime, self.SessionCode, self.RobotCode))
                            else:
                                pass
                    except:
                        pass
        except:
            print('link fail')
            pass

    def add_hyperlink(self, paragraph, text, url):

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        # hyperlink
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)
        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
        # print("Hyperlink type" + str(type(hyperlink)))

        return hyperlink

    def GetContent(self, url, selector, keyword, SessionCode):
        try:

            self.List_search.clear()
            self.Textcontent.clear()
            self.list_str.clear()
            req = requests.get(url)
            html = req.text
            # html = html.decode('cp1251')
            soup = BeautifulSoup(html, 'lxml')
            header_local = soup.find(selector["header"], class_=selector["headerclass"])
            header = header_local.text.replace('\n')
            header = header.replace('\t')
            header = header.strip()
            # print(header)
            Content_local = soup.find(selector["content"], class_=selector["contentclass"])
            Content = Content_local.text.replace('\n')
            Content = Content.replace('\t')
            Content = Content.strip()

            List_count = {
                'Key': '',
                'Times': 1
            }
            linkpostdata = ''
            contentdata = ''
            headerdata = ''
            if any(x in Content for x in keyword):
                for i in keyword:
                    List_count['Key'] = i
                    List_count['Times'] = Content.count(i) + header.count(i)
                    if List_count['Times'] == 0:
                        print("not have time----")
                        continue
                    else:
                        print("have content")

                        op = json.dumps(List_count)
                        print(op)
                        self.List_search.append(op)

                for j in self.List_search:
                    data = json.loads(j)
                    print(data)
                    self.list_str.append(data)

                self.scancontenturl.append(url)
                print(self.list_str)
                # list_str2 = list_str.replace("'",'"')
                list_str1 = json.dumps(self.list_str)
                print(list_str1)
                Url = url
                scan_json = {
                    'url': Url,
                    'keyword': list_str1,
                    'Status': True
                }
                print(scan_json)
                linkpostdata = 'Link post: ' + url
                headerdata = header
                contentdata = Content
                self.url_scan.append(scan_json)

            if linkpostdata == '' and contentdata == '' and headerdata == '':
                Url = url
                scan_json = {
                    'url': Url,
                    'keyword': self.list_str,
                    'Status': False
                }
                self.url_scan.append(scan_json)
            if linkpostdata != '' and contentdata != '' and headerdata != '' and str(self.list_str) != '':
                self.content.append(linkpostdata)
                self.content.append("Keywords search:" + str(self.List_search))
                self.content.append(headerdata)
                self.content.append(contentdata)
                self.Textcontent.append(headerdata)
                self.Textcontent.append(contentdata)
                # check inserrt to DB
                KeyWord = json.dumps(keyword)
                # print(KeyWord)
                text = '\n'.join(map(str, self.Textcontent))
                print(text)
                urlscan_json = json.dumps(self.List_search)
                # print(urlscan_json)
                # CREATE PROCEDURE [dbo].[P_GET_COUNT_CRAWL_DATA_POST]
                data = {
                    'BotCode': self.botcode[0],
                    'SessionCode': SessionCode,
                    'LinkPost': url,
                    'TextContent': text.strip(),
                    'MediaCrawl': 'check',
                    'KeyWord': KeyWord,
                    'KeyWordJson': urlscan_json,
                    'CreatedBy': 'Huycntt_3i',
                }

                print(data)
                url_upload = "http://localhost:6002/PythonCrawler/InsertCrawlerCrawlerSessionContent"
                resp = requests.post(url_upload, data=data)
                if resp.ok:
                    print("Upload completed successfully!")
                    print("data insert from Post!")
                    print(resp.text)
                    # API luecne index :
                    # resp.text = {"ID":316,"Title":"Thêm mới thành công","Error":false,"Object":null,"Code":null}
                    IdData = json.loads(resp.text)
                    print(IdData['ID'])
                    data1 = {
                        'fileCode': IdData['ID'],
                        'content': text.strip(),
                        'pathIndex': r'D:\wwwroot\III.SWORK.VATCO 2.2 NEW\wwwroot\uploads\repository\INDEX_SEARCH DEMO',
                        # D:\wwwroot\III.SWORK.VATCO 2.2 NEW\wwwroot\uploads\repository\INDEX_SEARCH - Crawler
                        'UrlPost': url,
                        'IdentifierCode': SessionCode,
                    }
                    print(data1)
                    url_upload = 'http://localhost:6002/PythonCrawler/PythonIndexContent'
                    resp1 = requests.post(url_upload, data=data1)
                    if resp1.ok:
                        print("Upload completed successfully!")
                        print(resp1.text)

                    else:
                        print("Something went wrong!")


                else:
                    print("Something went wrong!")

            return self.content
        except:
            pass

    def upload_data(self, Document):
        timebegin = self.StartTime
        print(timebegin)
        start_time1 = timebegin.strftime("%H:%M:%S")
        end_time = datetime.now()
        print(end_time)
        end_time1 = end_time.strftime("%H:%M:%S")
        myUUID = uuid.uuid4()
        mystr = str(myUUID)
        print(self.sesioncode)

        filename = self.sesioncode[0] + '_' + mystr[0:5] + '.docx'

        file_path = r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\spiders\save/' + filename
        print(file_path)
        print(Document)
        Document.save(file_path)
        file_exists = exists(file_path)
        print(file_exists)
        filesize = os.path.getsize(file_path)
        print(filesize)
        # r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\save\url1.docx')
        url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
        # url_upload = resp_json['DataStoragePath']

        response_upload = requests.post(url_upload, data={"CateRepoSettingId": 2247, "CreatedBy": "huynv_cntt_3i"},
                                        files={
                                            "fileUpload": (
                                                filename,
                                                open(
                                                    file_path,
                                                    'rb'),
                                                'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
        if response_upload.ok:
            print("Upload completed successfully!")
            print("Tải file lên DB!")
            print(response_upload.text)
            sleep(4)
            os.remove(file_path)
            statusupload = response_upload.text
            # {"ID":22577,"Title":"Thêm mới tệp tin thành công","Error":false,"Object":{"Type":"SERVER","Url":"https://dieuhanh.vatco.vn//uploads/repository/CRAWLING/vneconomy.docx","CloudId":""},"Code":null}
            datasend = json.loads(statusupload)
            File_Result = datasend['Object']["Url"]
        else:
            print("Something went wrong!")
            File_Result = ''
        # File_Result= file_path

        urlscan_json = json.dumps(self.url_scan)
        print(urlscan_json)
        # data["UrlScanJson"] = urlscan_json
        # data["endtime"] = end_time
        print(self.Lstkeyword)
        data = {
            'SessionCode': self.sesioncode[0],
            'StartTime': start_time1,
            'EndTime': end_time1,
            'UrlScanJson': urlscan_json,
            'FileDownloadJson': '',
            'NumOfFile': 1,
            'FileResultData': File_Result,
            'NumPasscap': '',
            'UserIdRunning': '001',
            'Ip': '1',
            'Status': 'active',
            'BotCode': self.botcode[0],
            'TimeScan': int((end_time - timebegin).total_seconds()),
            'CreatedBy': 'Huycntt_3i',
            'KeyWord': str(self.Lstkeyword),
            'FileSizeDownload': int(filesize),
        }
        # [FILE_SIZE_DOWNLOAD
        print(data)
        url_upload = "https://dieuhanh.vatco.vn/PythonCrawler/InsertCrawlerRunningLog"
        # url_upload = "https://dieuhanh.vatco.vn/PythonCrawler/InsertCrawlerRunningLog"
        resp = requests.post(url_upload, data=data)
        if resp.ok:
            print("Upload completed successfully!")
            print("data report crawl!")
            print(resp.text)
        else:
            print("Something went wrong!")

    async def main(self):

        Idx = 0
        len_list = len(self.list_child_url)
        await asyncio.sleep(10)
        while Idx < len_list:
            async with websockets.connect(URL, ping_interval=None) as ws:
                try:
                    self.Extract_Url(self.list_child_url[Idx].url, self.list_child_url[Idx].deep,
                                     self.list_child_url[Idx].deepscan, self.list_child_url[Idx].linkadd)
                except:
                    pass
                try:
                    self.GetContent(self.list_child_url[Idx].url, self.list_child_url[Idx].selector,
                                    self.list_child_url[Idx].keyword, self.list_child_url[Idx].SessionCode)
                    # print(list_child_url[Idx].url)
                    print("start search content")
                except:
                    pass
                await ws.send(str(datetime.now()) + ": " + self.list_child_url[Idx].url)
                Idx = Idx + 1
                len_list = len(self.list_child_url)
                pass
        document = Document()
        document.add_paragraph('Search With Keyword:' + str(self.Lstkeyword))
        print('Done. Start save text!')
        for value in self.content:
            p = document.add_paragraph(value)
            for url in self.scancontenturl:

                if url in value:
                    self.add_hyperlink(p, 'Link!', url)

                    # self.delete_paragraph(p)
        self.upload_data(document)
        # document.save(r'F:\PycharmProjects\Source\DomainManagement\DomainManagement\spiders\save/' + 'test.docx')
        # print(1)

        self.content.clear()
        self.list_child_url.clear()
        self.url_scan.clear()
        self.scancontenturl.clear()
        self.Lstkeyword.clear()
        self.botcode.clear()
        self.list_str.clear()
        self.sesioncode.clear()
        async with websockets.connect(URL, ping_interval=None) as ws:
            await ws.send('Write file and Finish!')
