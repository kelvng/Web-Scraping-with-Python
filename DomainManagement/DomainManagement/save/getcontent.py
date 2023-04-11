from bs4 import BeautifulSoup
import requests
import json
from docx import Document

name ="Geturl"


url = "https://vneconomy.vn/"

# list_child_url = []
content = []
check = []
class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep
    Idx = 0
    chk = 0
    with open(r'/DomainManagement/DomainManagement/spiders/selector5.json', 'r') as j:
            list_tag = json.loads(j.read())

    def GetContent(self, url):
        list_tag = self.list_tag
        req = requests.get(url)
        html = req.text
        soup = BeautifulSoup(html, 'html.parser')
        check = []
        check.clear()
        LstLink = soup.find('html')
        myListString = list_tag['string']
        for s in LstLink.find_all(list_tag['title']):
            try:
                title = s.text

                for x in myListString:
                    if x in title:
                        content.append(title)
                        check.append(title)
            except:
                pass
        if len(check) != 0:
            content.append(url)
        for s in LstLink.find_all(list_tag['summary'], class_=list_tag['summaryclass']):
            try:
                text = s.text
                for x in myListString:
                    if x in title:
                        content.append(text)
                        check.append(text)

            except:
                pass

        for img in LstLink.findAll(True, list_tag['imageclass']):
            image = img.find('img')['src']
            if "http" in image:
                content.append(image)
        # time = LstLink.find_all(list_tag['time'], class_= list_tag['timeclass'])
        # content.append(time)
        # for img in LstLink.find_all(list_tag['img'], class_=list_tag['imgclass']):
        #     image = img.find('img')['src']
        #     print(image)

            # with open("link_image", "w+b") as im_h:
            #     im_h.write(link_image.content)
        # if '' != title:
        #     content.append({
        #         title,
        #         text,
        #         # time,                  # image
        #         })
        return content

# list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
# object.main()
document = Document()
object.GetContent(url)

#
# a_keep_first = list(df.drop_duplicates(keep="first")['a'])
for value in content:

    document.add_paragraph(value)
    #document.add_picture('img_head.png')
    print(value)
print(len(content))
# Vị trí lưu file

#filePath = r'C:\Users\Admin\PycharmProjects\Source\DomainManagement\DomainManagement\save/'
document.save('C:\\Users\\Admin\\PycharmProjects\\Source\\DomainManagement\\DomainManagement\\save\\url.docx')