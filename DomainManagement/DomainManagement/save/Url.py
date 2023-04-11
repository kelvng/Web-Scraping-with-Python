import json
import sys
from hashlib import new

import requests
from bs4 import BeautifulSoup
from docx import Document
from scrapy.selector import Selector


def url_obj(url, param, param1):
    pass
#Tên spider
name ="Geturl"


# resp = requests.post("https://os.3i.com.vn/PythonCrawler/GetCrawlerData?spiderName=crawler")
# resp_json = resp.json()
# # load post tạo json
# print(resp_json)
# print(resp_json['Url'])
# url = resp_json['Url']
url = "https://en.vneconomy.vn/pre-feasibility-study-for-hcmcs-ring-road-no-3-to-go-to-na.htm"

date = []
Title = []
author = []
image = []
summary = []
list_child_url = []
content = []

#mảng chứa lisr url và content text

class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep
    Idx = 0
    chk = 0
    # Hàm kiểm tra link đã tồn tại trong list chưa
    def chk_link_exist(self, link):
        for obj in list_child_url:
            if (obj.url == link):
                 return 1
        return 0
    #Hàm quét và lấy các link con từ link mẹ
    def Extract_Url(self,url,deep):
        try:
            # Độ sâu của link mẹ là 1, link con = link mẹ +1
            if deep <= 2:
                req = requests.get(url)
                html = req.text
                soup = BeautifulSoup(html, 'html.parser')
                LstLink = soup.find('body')
                for s in LstLink.find_all('a'):
                    try:
                        link = s['href']
                        if 'https' not in link:
                            link = 'https://en.vneconomy.vn' + link
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        else:
                            if self.chk_link_exist(link) == 0:
                                url = link
                                iscan = 0
                                deep = self.deep + 1
                                list_child_url.append(url_obj(url, iscan, deep))
                        print(link)
                    except:
                        pass
        except:
            print('link fail')
            pass


    # with open('/DomainManagement/DomainManagement/spiders/selector2.json', 'r') as j:
    #         list_tag = json.loads(j.read())
    #
    # def GetContent(self, url):
    #     list_tag = self.list_tag
    #     req = requests.get(url)
    #     html = req.text
    #     soup = BeautifulSoup(html, 'html.parser')
    #     LstLink = soup.find('html')
    #     for s in LstLink.find_all(list_tag['title']):
    #         title = s.text
    #         if '' != title:
    #             content.append(title)
    #
    #     return content

    def main(self):
        Idx = 0
        len_list = len(list_child_url)
        while Idx < len_list:
            self.Extract_Url(list_child_url[Idx].url,list_child_url[Idx].deep)
            Idx = Idx + 1
            len_list = len(list_child_url)

document = Document()
list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
object.main()
# document = Document()
# object.GetContent(url)
for obj in list_child_url:
    document.add_paragraph(obj.url)
    print(obj.url, obj.iscan, obj.deep, sep=' ')
print(len(list_child_url))
for value in content:
    document.add_paragraph(value)
    print(value)
# for sum in summary:
#     document.add_paragraph(sum)
#     print(sum)
print(len(content))
# Vị trí lưu file
filePath = r'\DomainManagement\DomainManagement\spiders\url.docx'
document.save("C:\\Users\\Admin\\PycharmProjects\\Source" + filePath)


sys.exit()