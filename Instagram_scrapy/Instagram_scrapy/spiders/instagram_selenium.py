# coding=utf8
# -*- coding: utf-8 -*-
'''https://pypi.org/project/instagramy/'''
import asyncio
import glob
import json
import os
import random
import sys
from datetime import datetime
from time import sleep
import requests
import scrapy
import websockets
from docx import Document
from instagramy import InstagramHashTag
from pywinauto import Application
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.keys import Keys

options = webdriver.EdgeOptions()
prefs = {"profile.default_content_setting_values.notifications": 2}
options.add_experimental_option("prefs", prefs)
options.add_argument("start-maximized")
driver_path = r'C:\Users\Admin 3i\PycharmProjects\Source\msedgedriver.exe'
browser = webdriver.Edge(executable_path=driver_path, options=options)

URL = 'ws://127.0.0.1:9091'

def wait():
    return sleep(random.randint(4, 6))

def Bot_func():
    async def login(TK, MK):
        async with websockets.connect(URL, ping_interval=None) as ws:
            browser.get("https://www.instagram.com/accounts/login/")
            sleep(1)
            # login to instagram
            id = browser.find_element_by_xpath('//*[@aria-label="Phone number, username, or email"]')
            id.send_keys(TK)
            sleep(2)
            pwd = browser.find_element_by_xpath('//*[@aria-label="Password"]')
            pwd.send_keys(MK)
            pwd.send_keys(Keys.ENTER)
            await ws.send('Login: ' + TK)
            sleep(1)
            try:
                browser.find_element_by_xpath("//button[contains(text(),'Save Info')]").click()
            except NoSuchElementException:
                print("no save Info")
            try:
                browser.find_element_by_xpath("//*[contains(@class, 'aOOlW   HoLwm ')]").click()
            except NoSuchElementException:
                print("no notification box")

            sleep(2)

    async def instagram_auto_post(media, content, friend):
        async with websockets.connect(URL, ping_interval=None) as ws:
            await ws.send('Start post')
            try:
                browser.find_element_by_xpath('//*[@aria-label="New Post"]').click()
            except:
                browser.refresh()
                browser.find_element_by_xpath('//*[@aria-label="New Post"]').click()
            # sleep(2)
            lst_img = []
            image_path = r"C:\Users\Admin 3i\PycharmProjects\Source\Instagram_scrapy\Instagram_scrapy\spiders\saveimg"
            for i in range(len(media)):
                image = requests.get(media[i]).content
                name = 'img' + str(i) + '.jpg'
                with open(image_path + '/' + name, 'wb') as f:
                    f.write(image)
                lst_img.append(name)
            file = str(" ".join('"{}"'.format(i) for i in lst_img))
            print(file)
            # "img4.jpg" "img1.jpg" "img2.jpg" "img3.jpg"

            browser.find_element_by_xpath(
                '//button[contains(text(), "Select from computer")]').click()
            sleep(2)
            await asyncio.sleep(3)
            upload_dialog = Application().connect(title_re='Open')
            upload_dialog.Open.Edit.type_keys(image_path)
            sleep(2)
            upload_dialog.window(title_re='Open').Open.click()
            sleep(1)
            upload_dialog.Open.Edit.type_keys(file)
            await asyncio.sleep(3)
            upload_dialog.window(title_re='Open').Open.click()
            sleep(1)
            browser.find_element_by_xpath(
                '//button[contains(text(), "Next")]').click()
            wait()
            browser.find_element_by_xpath(
                '//button[contains(text(), "Next")]').click()
            wait()
            caption = browser.find_element_by_xpath('//*[@aria-label="Write a caption..."]')
            caption.send_keys(content + '\n' + '@' + friend)
            browser.find_element_by_xpath(
                '//button[contains(text(), "Share")]').click()
            await ws.send('Posted: ' + content)
            filelist = glob.glob(os.path.join(image_path, "*"))
            for g in filelist:
                os.remove(g)

    async def instagram_auto_comment(content, hashtag, from_time, to_time):
        for i in range(len(hashtag)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Start commend!')
                session_id = os.environ.get("51620542060%3AyyD9gGUI0VAxQz%3A19")
                tag = InstagramHashTag(hashtag[i], sessionid=session_id)
                data_string = json.dumps(tag.tag_data, indent=4)
                data = json.loads(data_string)
                edges = data['edge_hashtag_to_media']['edges']

                print(len(edges))
                # link post hashtag
                link = []
                for j in range(len(edges)):
                    time_stamp = edges[j]['node']['taken_at_timestamp']
                    dt_obj = datetime.fromtimestamp(time_stamp)
                    if check_time(from_time, to_time, dt_obj) == 1:
                        id_post = edges[j]['node']['shortcode']
                        text_post = edges[j]['node']['edge_media_to_caption']['edges'][0]['node']['text']
                        print(text_post)
                        link_post = 'https://www.instagram.com/p/' + str(id_post)
                        print(link_post)
                        link.append(link_post)
                for i in link_post:
                    async with websockets.connect(URL, ping_interval=None) as ws:
                        await ws.send('Start commend!')
                        browser.get(i)
                        sleep(2)
                        browser.find_element_by_xpath('//*[@class="wG4fl  UDpcu "]').click()
                        comment_box = browser.find_element_by_xpath('//*[@data-testid="post-comment-text-area"]')
                        comment_box.send_keys(content)
                        sleep(2)
                        browser.find_element_by_xpath('//*[@data-testid="post-comment-input-button"]').click()
                        sleep(1)

    async def instagram_auto_like(postID, hashtag, from_time, to_time):
        try:
            for i in range(len(hashtag)):
                session_id = os.environ.get("51620542060%3AyyD9gGUI0VAxQz%3A19")
                tag = InstagramHashTag(hashtag[i], sessionid=session_id)
                data_string = json.dumps(tag.tag_data, indent=4)
                data = json.loads(data_string)
                edges = data['edge_hashtag_to_media']['edges']
                print(len(edges))
                # link post hashtag
                link = []
                for j in range(len(edges)):
                    time_stamp = edges[j]['node']['taken_at_timestamp']
                    dt_obj = datetime.fromtimestamp(time_stamp)
                    if check_time(from_time, to_time, dt_obj) == 1:
                        id_post = edges[j]['node']['shortcode']
                        link_post = 'https://www.instagram.com/p/' + str(id_post)
                        link.append(link_post)
            for g in range(len(link)):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send('Start like!')
                    browser.get(link[g])
                    # '''tìm ô like theo nhãn, tuỳ theo ngôn ngữ mà sẽ khác nhau. Việc tìm nhãn phải làm manual'''
                    try:
                        sleep(2)
                        browser.find_element_by_xpath('//*[@class="QBdPU rrUvL"]').click()
                        await ws.send('Liked: ' + link[g])
                    except:
                        sleep(1)
        except:
            pass
        try:
            for g in range(len(postID)):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send('Start like!')
                    browser.get('https://www.instagram.com/p/' + postID[g])
                    # '''tìm ô like theo nhãn, tuỳ theo ngôn ngữ mà sẽ khác nhau. Việc tìm nhãn phải làm manual'''
                    try:
                        sleep(2)
                        browser.find_element_by_xpath('//*[@class="QBdPU rrUvL"]').click()
                        await ws.send('Liked: ' + 'https://www.instagram.com/p/' + postID[g])
                    except:
                        sleep(1)
        except:
            pass

    async def instagram_auto_unlike(postID, hashtag, keyword, from_time, to_time):
        try:
            for i in range(len(hashtag)):
                session_id = os.environ.get("51620542060%3AyyD9gGUI0VAxQz%3A19")
                tag = InstagramHashTag(hashtag[i], sessionid=session_id)
                data_string = json.dumps(tag.tag_data, indent=4)
                data = json.loads(data_string)
                edges = data['edge_hashtag_to_media']['edges']
                print(len(edges))
                # link post hashtag
                link = []
                for j in range(len(edges)):
                    time_stamp = edges[j]['node']['taken_at_timestamp']
                    dt_obj = datetime.fromtimestamp(time_stamp)
                    if check_time(from_time, to_time, dt_obj) == 1:
                        id_post = edges[j]['node']['shortcode']
                        link_post = 'https://www.instagram.com/p/' + str(id_post)
                        print(link_post)
                        link.append(link_post)
                for g in range(len(link)):
                    async with websockets.connect(URL, ping_interval=None) as ws:
                        await ws.send('Start Unlike!')
                        browser.get(link[g])
                        # '''tìm ô like theo nhãn, tuỳ theo ngôn ngữ mà sẽ khác nhau. Việc tìm nhãn phải làm manual'''
                        try:
                            wait()
                            browser.find_element_by_xpath('//*[@class="fr66n"]').click()
                            await ws.send('Unliked: ' + link[g])
                        except:
                            wait()
        except:
            pass
        try:
            for g in range(postID):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send('Start Unlike!')
                    browser.get('https://www.instagram.com/p/' + postID[g])
                    # '''tìm ô like theo nhãn, tuỳ theo ngôn ngữ mà sẽ khác nhau. Việc tìm nhãn phải làm manual'''
                    try:
                        wait()
                        browser.find_element_by_xpath('//*[@class="fr66n"]').click()
                        await ws.send('Unliked: ' + 'https://www.instagram.com/p/' + postID[g])
                    except:
                        wait()
        except:
            pass

    def getlink_hashtag(hashtag, keyword, from_time, to_time):
        for i in range(len(hashtag)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Start Search!')
                await ws.send('Search Hashtag: ' + str(hashtag))

                session_id = os.environ.get("51620542060%3AyyD9gGUI0VAxQz%3A19")
                tag = InstagramHashTag(hashtag[i], sessionid=session_id)
                data_string = json.dumps(tag.tag_data, indent=4)
                data = json.loads(data_string)
                edges = data['edge_hashtag_to_media']['edges']
                print(len(edges))
                # link post hashtag
                link = []
                for j in range(len(edges)):
                    id_post = edges[j]['node']['shortcode']
                    link_post = 'https://www.instagram.com/p/' + str(id_post)
                    print(link_post)
                    link.append(link_post)
                return link

    async def instagram_auto_follow(ID, hashtag):
        if len(ID) != 0:
            for i in range(len(ID)):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send('Start follow!')
                    browser.get('https://www.instagram.com/' + str(ID[i]))
                    try:
                        sleep(2)
                        browser.find_element_by_xpath('//*[@class="vBF20 _1OSdk"]').click()
                        await ws.send('Followed: ' + str(ID[i]))
                    except:
                        await ws.send('No username: ' + str(ID[i]) + 'searching')
                        browser.get('https://www.instagram.com/')
                        sleep(1)
                        box = browser.find_element_by_xpath('//*[@aria-label="Phần nhập nội dung tìm kiếm"]')
                        box.send_keys(ID[i])
                        sleep(1)
                        box.send_keys(Keys.ENTER)
                        box.send_keys(Keys.ENTER)
                        sleep(2)
                        browser.find_element_by_xpath('//*[@class="vBF20 _1OSdk"]').click()
                        name = browser.find_element_by_xpath('//*[@class="_7UhW9       fKFbl yUEEX    KV-D4               fDxYl     "]').text
                        await ws.send('Followed: ' + name)
                        pass
        else:
            pass
        if len(hashtag) != 0:
            for j in range(len(hashtag)):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send('Start follow!')
                    browser.get('https://www.instagram.com/explore/tags/' + hashtag[j])
                    try:
                        browser.find_element_by_xpath('//*[@class="sqdOP  L3NKy _4pI4F  y3zKF     "]').click()
                        await ws.send('Followed: '+ '#' + hashtag[j])
                    except:
                        await ws.send('No Hashtag or Hashtag cannot be followed!')
                        pass

    async def instagram_auto_unfollow(ID):
        for i in range(len(ID)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Start unfollow!')
                browser.get('https://www.instagram.com/' + str(ID[i]))
                try:
                    browser.find_element_by_xpath('//*[@class="vBF20 _1OSdk"]').click()
                    sleep(2)
                    browser.find_element_by_css_selector('.-Cab_').click()
                    await ws.send('Unfollowed: ' + 'https://www.instagram.com/' + str(ID[i]))
                except:
                    sleep(1)
    async def instagram_auto_search(hashtag, from_time, to_time):
        for i in range(len(hashtag)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Start Search!')
                await ws.send('Search Hashtag: ' + str(hashtag))
                content = []
                session_id = os.environ.get("51620542060%3AyyD9gGUI0VAxQz%3A19")
                tag = InstagramHashTag(hashtag[i], sessionid=session_id)
                data_string = json.dumps(tag.tag_data, indent=4)
                data = json.loads(data_string)
                edges = data['edge_hashtag_to_media']['edges']
                print(len(edges))
                # link post hashtag
                for j in range(len(edges)):
                    id_post = edges[j]['node']['shortcode']
                    time_stamp = edges[j]['node']['taken_at_timestamp']
                    dt_obj = datetime.fromtimestamp(time_stamp)
                    text_post = edges[j]['node']['edge_media_to_caption']['edges'][0]['node']['text']
                    link_post = 'https://www.instagram.com/p/' + str(id_post)
                    if check_time(from_time,to_time,dt_obj) == 1:
                        content.append(link_post)
                        content.append(text_post)
                        content.append(str(dt_obj))

                document = Document()
                endtime = datetime.now()
                end_time = endtime.strftime("%H:%M:%S")
                end_time = end_time.replace(":", "_")
                filename = 'INSearch' + hashtag[i] + end_time
                for value in content:
                    document.add_paragraph(value)
                await ws.send('Get text of: ' + str(hashtag) + ' Done')
                document.save(r'C:\Users\Admin 3i\PycharmProjects\Source\Instagram_scrapy/' + filename + '.docx')
                url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
                response_upload = requests.post(url_upload,
                                                data={"CateRepoSettingId": 2247, "CreatedBy": "phancuoc_cntt_3i"},
                                                files={
                                                    "fileUpload": (
                                                        filename + '.docx',
                                                        open(
                                                            r'C:\Users\Admin 3i\PycharmProjects\Source\Instagram_scrapy/' + filename + '.docx',
                                                            'rb'),
                                                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
                if response_upload.ok:
                    print("Upload completed successfully!")
                    print(response_upload.text)
                else:
                    print("Something went wrong!")
        # BotSessionResult = {
        #     'Numkeyword': len(keyword),
        #     'CountLike': dem
        # }
        # # document = Document()
        endtime = datetime.now()
        end_time = endtime.strftime("%H:%M:%S")
        end_time = end_time.replace(":", "_")
        filename = None
        BotSessionResult = json.dumps('BotSessionResult')
        asyncio.create_task(post_log(filename, BotSessionResult))
    def check_keyword(keyword, content):
        try:
            for x in keyword:
                if x in content:
                    return 1
            return 0
        except:
            pass

    def check_time(from_time, to_time, time_post):
        from_time1 = from_time.replace('T', ' ')
        from_time2 = from_time1.replace('.000Z', '')
        to_time1 = to_time.replace('T', ' ')
        to_time2 = to_time1.replace('.000Z', '')
        if datetime.strptime(str(time_post), "%Y-%m-%d %H:%M:%S") > datetime.strptime(from_time2,
                                                                                      "%Y-%m-%d %H:%M:%S") and datetime.strptime(
            str(time_post), "%Y-%m-%d %H:%M:%S") < datetime.strptime(to_time2, "%Y-%m-%d %H:%M:%S"):
            return 1
        else:
            return 0

    async def init_unfollow(ID):
        print("unfollow")
        await asyncio.sleep(5)
        asyncio.create_task(instagram_auto_unfollow(ID))

    async def post_log(filename, result):
        endtime = datetime.now()
        end_time = endtime.strftime("%H:%M:%S")
        end_time = end_time.replace(":", "_")

        data = {
            'SessionCode': 'Botname[0]' + str(end_time),
            'StartTime': 'start_time',
            'EndTime': datetime.now(),
            'Statvs': '',
            'BotSocialCode': 'Botid[0]',
            'FileResults': filename,
            'RuningType': 'facebook',
            'BotSessionResult': result,
            'CreatedBy': 'admin',
        }
        data = json.dumps(data)
        url_upload = "http://localhost:6023/PythonCrawler/InsertBotSocialLog"
        resp = requests.post(url_upload, data=data)
        if resp.ok:
            print("Upload completed successfully!")
            print(resp.text)
        else:
            print("Something went wrong!")

    async def listen():
        ws_connect = websockets.connect('ws://127.0.0.1:9091', ping_interval=None)
        async with ws_connect as wb:
            await wb.send('Spider running!')
            while True:
                param = await wb.recv()

                if "UserName" in param:
                    data = json.loads(param)
                    tk = data['UserName']
                    mk = data['PassWord']
                    asyncio.create_task(login(tk, mk))
                if "Start post content" in param:
                    data = json.loads(param)
                    media = data['Start post content']['media']
                    content = data['Start post content']['content']
                    friend = data['Start post content']['friends']
                    await asyncio.sleep(3)
                    asyncio.create_task(instagram_auto_post(media, content, friend))

                if "Start post comment" in param:
                    data = json.loads(param)
                    linkpost = data['Start post comment']['linkpost']
                    content = data['Start post comment']['content']
                    friend = data['Start post comment']['friends']
                    hashtag = ["Trump", "Education", "life"]
                    from_time = data['Start post comment']['from']
                    to_time = data['Start post comment']['to']
                    await asyncio.sleep(5)
                    asyncio.create_task(instagram_auto_comment(content, hashtag, from_time, to_time))

                if "userid" in param:
                    data = json.loads(param)
                    id = data['userid']
                    hashtag= data['hahstag']
                    await asyncio.sleep(5)
                    asyncio.create_task(instagram_auto_follow(id, hashtag))
                if "Unfollow" in param:
                    data = json.loads(param)
                    id = data['Unfollow']['ID']
                    await asyncio.sleep(5)
                    asyncio.create_task(init_unfollow(id))
                if "Like post" in param:
                    data = json.loads(param)
                    postid = data['Like post']['postId']
                    hashtag = ["Trump","Education","life"]
                    from_time = data['Like post']['from']
                    to_time = data['Like post']['to']
                    await asyncio.sleep(5)
                    asyncio.create_task(instagram_auto_like(postid,hashtag, from_time, to_time))

                if "unlikepost" in param:
                    data = json.loads(param)
                    postid = data['unlikepost']['post']
                    keyword = 2
                    time = 3
                    await asyncio.sleep(5)
                    asyncio.create_task(instagram_auto_unlike(postid, hashtag, keyword, time))

                if "gethashtag" in param:
                    data = json.loads(param)
                    hashtag = data['gethashtag']['Hashtag']
                    from_time = data['gethashtag']['from']
                    to_time = data['gethashtag']['to']
                    await asyncio.sleep(5)
                    asyncio.create_task(instagram_auto_search(hashtag, from_time, to_time))

                if 'Stop post' in param:
                    browser.quit()
                    sys.exit()

    async def main():
        task_1 = asyncio.create_task(listen())
        await asyncio.sleep(0.250)
        await task_1

    async def forever():
        while True:
            await main()

    loop = asyncio.get_event_loop()
    loop.run_until_complete(forever())
class Myspider(scrapy.Spider):
    name = 'CoreInstagramSpider'
    some_attribute = "Yes|No"
    print('Hello world!')
    Bot_func()

# if __name__ == '__main__':
#     asyncio.run(main())
