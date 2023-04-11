# coding=utf8
# -*- coding: utf-8 -*-
import asyncio
import json
import random
import re
import sys
from datetime import datetime
from time import sleep
import pandas as pd
import requests
import websockets
from bs4 import BeautifulSoup
from docx import Document
from facebook_scraper import get_posts
from pynput.keyboard import Controller
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support import expected_conditions as ec
from selenium.webdriver.support.ui import WebDriverWait
from tqdm import tqdm
import scrapy

for post in get_posts('nintendo', pages=1):
    print(post['text'][:50])

options = webdriver.ChromeOptions()
prefs = {"profile.default_content_setting_values.notifications": 2}
options.add_experimental_option("prefs", prefs)
options.add_argument("start-maximized")
browser = webdriver.Chrome(executable_path=r'C:\Users\Admin 3i\PycharmProjects\Source\chromedriver.exe',
                           options=options)

def wait():
    return sleep(random.randint(5, 8))


start_time = datetime.now()
lstcookie = []
Botresult = []
Botname = []
Botid = []
URL = 'ws://127.0.0.1:9091'
def Bot_func():
    async def loginFacebookByCookie(cookie):
        async with websockets.connect(URL, ping_interval=None) as ws:
            await ws.send('Start Login!')
            try:
                browser.get('https://facebook.com/')
                sleep(1)
            except:
                print("check live fail")
            try:
                new_cookie = ["c_user=", "xs="]
                cookie_arr = cookie.split(";")
                for i in cookie_arr:
                    if i.__contains__('c_user='):
                        new_cookie[0] = new_cookie[0] + (i.strip() + ";").split("c_user=")[1]
                    if i.__contains__('xs='):
                        new_cookie[1] = new_cookie[1] + (i.strip() + ";").split("xs=")[1]
                        if (len(new_cookie[1].split("|"))):
                            new_cookie[1] = new_cookie[1].split("|")[0]
                        if (";" not in new_cookie[1]):
                            new_cookie[1] = new_cookie[1] + ";"
                conv = new_cookie[0] + " " + new_cookie[1]
            except:
                print("Error Convert Cookie")
            try:
                cookie = conv
                #print(cookie)
                if (cookie != None):
                    script = 'javascript:void(function(){ function setCookie(t) { var list = t.split("; ");' \
                             ' console.log(list); for (var i = list.length - 1; i >= 0; i--) { var cname = list[i].split("=")[0];' \
                             ' var cvalue = list[i].split("=")[1]; var d = new Date(); d.setTime(d.getTime() + (7*24*60*60*1000));' \
                             ' var expires = ";domain=.facebook.com;expires="+ d.toUTCString(); document.cookie = cname + "=" + cvalue + "; " + expires; } }' \
                             ' function hex2a(hex) { var str = ""; for (var i = 0; i < hex.length; i += 2) { var v = parseInt(hex.substr(i, 2), 16); if (v) str += String.fromCharCode(v); }' \
                             ' return str; } setCookie("' + cookie + '"); location.href = "https://facebook.com"; })();'

                    # script1 = ' javascript:void(function(){ function setCookie(t) { var list = t.split("; "); console.log(list); for (var i = list.length - 1; i >= 0; i--) { var cname = list[i].split("=")[0]; var cvalue = list[i].split("=")[1]; var d = new Date(); d.setTime(d.getTime() + (7*24*60*60*1000)); var expires = ";domain=.facebook.com;expires="+ d.toUTCString(); document.cookie = cname + "=" + cvalue + "; " + expires; } } function hex2a(hex) { var str = ""; for (var i = 0; i < hex.length; i += 2) { var v = parseInt(hex.substr(i, 2), 16); if (v) str += String.fromCharCode(v); } return str; } setCookie("{c_user=100073734591574; xs=42%3AGfemPI2zWBeNMw%3A2%3A1677121112%3A-1%3A6279%3A%3AAcWM56DE-oBzs7y1Lgy6Fuc6LkBog938Wvw5dgBjSQ;}"); location.href = "https://facebook.com"; })();'
                    browser.execute_script(script)
                    await ws.send('Login success!')
                    print('Login success!')
            except:
                print("Error login")
                await ws.send('Login Fail pls check tocken!')
            lstcookie.append(cookie)

    async def facebook_auto_post_group(group, content):
        if len(group) != 0:
            for i in group:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send('Start Post group!')
                    await  asyncio.sleep(5)
                    url = 'https://www.facebook.com/groups/' + str(i['id'])
                    browser.get(url)
                    WebDriverWait(browser, 20).until(EC.visibility_of_element_located((By.XPATH,
                    '//*[@class="x1lkfr7t xkjl1po x1mzt3pk xh8yej3 x13faqbe xi81zsa"]'))).click()

                    try:
                        WebDriverWait(browser, 10).until(ec.visibility_of_element_located(
                            (By.XPATH,
                             '//*[@class="x76ihet xwmqs3e x112ta8 xxxdfa6 x9f619 xzsf02u xmper1u xo1l8bm x5yr21d x1a2a7pz x1iorvi4 x4uap5 xwib8y2 xkhd6sd xh8yej3 xha3pab xngnso2 x1qb5hxa"]')))
                        bd = browser.find_element(By.XPATH,
                                  '//*[@aria-label="Tạo bài viết công khai..."]')
                        bd.send_keys(content['content'])
                        bd.send_keys(Keys.ENTER)
                        bd.send_keys( content['Image'])
                        sleep(2)
                        await ws.send('Post with group :' + url)
                        WebDriverWait(browser, 10).until(ec.element_to_be_clickable(
                            (By.XPATH, '//*[@aria-label="Đăng"]'))).click()
                        # bt = browser.find_element_by_xpath('//*[@aria-label="Đăng"]')
                        WebDriverWait(browser, 10).until(
                            ec.invisibility_of_element((By.XPATH, '//*[@method="POST"]')))
                        sleep(3)
                        await ws.send('Done!')

                    except:
                        pass
                    try:
                        WebDriverWait(browser, 10).until(ec.visibility_of_element_located(
                            (By.XPATH, '//*[@aria-label="Bạn viết gì đi..."]')))
                        bd = browser.find_element_by_xpath('//*[@aria-label="Bạn viết gì đi..."]')
                        bd.send_keys(content['content'])
                        bd.send_keys('\n' + content['Image'])
                        sleep(2)
                        await ws.send('Post with group :' + url)
                        WebDriverWait(browser, 10).until(ec.element_to_be_clickable(
                            (By.XPATH, '//*[@aria-label="Đăng"]'))).click()
                        # bt = browser.find_element_by_xpath('//*[@aria-label="Đăng"]')
                        WebDriverWait(browser, 10).until(
                            ec.invisibility_of_element((By.XPATH, '//*[@method="POST"]')))
                        sleep(3)
                        await ws.send('Done!')

                    except:
                        pass
        if len(group) == 0:
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send('Start Post!')
                await  asyncio.sleep(7)
                url = 'https://www.facebook.com'
                browser.get(url)
                await ws.send('Post to profile!')
                WebDriverWait(browser, 20).until(EC.visibility_of_element_located((By.XPATH,
               '/html/body/div[1]/div/div[1]/div/div[3]/div/div/div/div[1]/div[1]/div/div[2]/div/div/div/div[3]/div/div[2]/div/div/div/div[1]/div/div[1]'))).click()

                sleep(1)
                # check selector of post main
                bd = WebDriverWait(browser, 20).until(EC.visibility_of_element_located((By.XPATH,
                '/html/body/div[1]/div/div[1]/div/div[4]/div/div/div[1]/div/div[2]/div/div/div/form/div/div[1]/div/div/div/div[2]/div[1]/div[1]/div[1]/div/div/div[1]/p')))

                # bd = browser.find_element(By.XPATH,'//*[@id="mount_0_0_Vt"]/div/div[1]/div/div[4]/div/div/div[1]/div/div[2]/div/div/div/form/div/div[1]/div/div/div/div[2]/div[1]/div[1]/div[1]/div/div/div[1]')
                bd.send_keys(content['content'])
                bd.send_keys(Keys.ENTER)
                bd.send_keys(content['Image'])
                await ws.send('Post content :' + content['content'] + '\n' + content['Image'])
                # sleep(2)
                WebDriverWait(browser, 10).until(ec.element_to_be_clickable(
                    (By.XPATH, '//*[@aria-label="Đăng"]'))).click()
                # bt = browser.find_element_by_xpath('//*[@aria-label="Đăng"]')
                sleep(3)
                await ws.send('Done!')
                print('Hoàn thành')
    # done json fixx comment in linkpost

    # async def facebook_auto_comment(linkpost, content):
    #     for i in range(len(linkpost)):
    #         async with websockets.connect(URL, ping_interval=None) as ws:
    #             await ws.send('Start Post comment!')
    #             browser.get(linkpost[i])
    #             await ws.send('link post: ' + linkpost[i])
    #             try:
    #                 await asyncio.sleep(4)
    #                 browser.find_element_by_xpath('//*[@aria-label="Viết bình luận"]').click()
    #
    #                 await asyncio.sleep(2)
    #                 comment_box = browser.find_element_by_xpath(('//*[@class="hcukyx3x oygrvhab cxmmr5t8 kvgmc6g5"]'))
    #                 comment_box.send_keys(content['content'] + '\n')
    #                 comment_box.send_keys(content['Image'])
    #                 await ws.send("commend: " + content['content'] + '\n' + content['Image'])
    #                 await asyncio.sleep(3)
    #                 comment_box.send_keys(Keys.ENTER)
    #                 sleep(1)
    #             except:
    #                 pass
    async def facebook_auto_comment(group, content, keyword, from_time, to_time):
        list_id = []
        dem = 0
        # get post from group
        for i in range(len(group)):
            data = get_posts(group=group[i]['id'],
                             cookie=r'H:\PycharmProjects\Source\facebook_scrapy\facebook_scrapy\spiders\cookie.json',
                             pages=2)

            data = pd.DataFrame.from_dict(data)
            a = 0
            print(data)
            while a < len(data):
                time_post = data['time'][a]
                content1 = data['text'][a]
                User = data['username'][a]
                React = data['reactors'][a]
                print(time_post)
                print(User)
                #print(content1)
                if check_keyword(keyword, content1) == 1:
                    if check_time(from_time, to_time, time_post) == 1:
                        list_id.append(data['post_id'][a])
                a += 1
            print(len(list_id))
            for j in range(len(list_id)):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    print(list_id[j])
                    browser.get("https://www.facebook.com/" + list_id[j])
                    WebDriverWait(browser, 10).until(
                        ec.presence_of_element_located((By.XPATH, '//*[@aria-label="Viết bình luận"]')))
                    await ws.send('link post: ' + "https://www.facebook.com/" + list_id[j])
                    await asyncio.sleep(1)

                    #browser.find_element(By.XPATH, '//*[@aria-label="Viết bình luận"]').click()
                    comment_box = browser.find_element(By.XPATH,'//*[@aria-label="Viết bình luận"]')
                    comment_box.send_keys(content['content'])
                    print("typing")
                    sleep(1)
                    comment_box.send_keys(Keys.ENTER)
                    print("Send")
                    await ws.send("commend: " + content['content'] + '\n' + content['Image'])
                    dem += 1
                    await asyncio.sleep(1)

    async def get_post_content(group, keyword, from_time, to_time):
        Lst_content = []
        dem = 0
        for i in range(len(group)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send("Start get post content!")
                try:
                    data = get_posts(group=group[i]['id'],
                                     cookie=r'H:\PycharmProjects\Source\facebook_scrapy\facebook_scrapy\spiders\cookie.json',
                                     pages=2)
                except:
                    print("No more posts to get")
                data = pd.DataFrame.from_dict(data)
                try:
                    await ws.send('collect data group: ' + 'https://www.facebook.com/' + group[i]['id'])
                except:
                    pass
                a = 0
                while a < len(data):
                    time_post = data['time'][a]
                    content = data['text'][a]
                    username = data['username'][a]
                    postid = data['post_id'][a]
                    link = 'https://www.facebook.com/' + str(postid)
                    if check_keyword(keyword, content) == 1:
                        if check_time(from_time, to_time, time_post) == 1:
                            # if check_friend(username, friend) == 1:
                            print(1)
                            Lst_content.append(
                                str(datetime.now()) + ", post at fiend: " + username + '\n' + "Link post: " + link)
                            Lst_content.append('Content: ' + content)
                            try:
                                await ws.send(str(datetime.now()) + ", post of: " + username + '\n' + postid)
                            except:
                                pass
                            dem += 1
                    a += 1
        BotSessionResult = {
            'Numkeyword': len(keyword),
            'Numblinkget': dem
        }
        document = Document()
        endtime = datetime.now()
        end_time = endtime.strftime("%H:%M:%S")
        end_time = end_time.replace(":", "_")
        filename = 'GetContent' + str(end_time)
        print(filename)
        for value in Lst_content:
            document.add_paragraph(value)
        document.save(r'H:\PycharmProjects\Source\facebook_scrapy/' + filename + '.docx')
        url_upload = "https://dieuhanh.vatco.vn/MobileLogin/InsertFile"
        response_upload = requests.post(url_upload,
                                        data={"CateRepoSettingId": 2247, "CreatedBy": "phancuoc_cntt_3i"}, files={
                "fileUpload": (
                    filename + '.docx',
                    open(r'H:\PycharmProjects\Source\facebook_scrapy/' + filename + '.docx', 'rb'),
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
        if response_upload.ok:
            print("Upload completed successfully!")
            print(response_upload.text)
        else:
            print("Something went wrong!")

        BotSessionResult = json.dumps(BotSessionResult)
        asyncio.create_task(post_log(filename, BotSessionResult))

    async def facebook_auto_like(group, keyword, from_time, to_time):
        list_id = []
        dem = 0
        for i in range(len(group)):
            data = get_posts(group=group[i]['id'],
                             cookie=r'C:\Users\Admin 3i\PycharmProjects\Source\facebook_scrapy\facebook_scrapy\spiders\cookie.json',
                             pages=2)
            data = pd.DataFrame.from_dict(data)
            a = 0
            # Quét id_post từ group
            while a < len(data):
                time_post = data['time'][a]
                content = data['text'][a]
                User = data['username'][a]
                print(User)
                if check_keyword(keyword, content) == 1:
                    if check_time(from_time, to_time, time_post) == 1:
                        list_id.append(data['post_id'][a])
                a += 1
            print(len(list_id))
            for product in list_id:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    url = "https://www.facebook.com/" + str(product)
                    browser.get(url)
                    try:
                        await ws.send("like post: " + url)
                    except:
                        pass
                    sleep(3)
                    try:
                        browser.find_element_by_xpath('//*[@aria-label="Thích"]').click()
                        await ws.send('Like post at ' + str(datetime.now()))
                        dem += 1
                    except:
                        pass

        BotSessionResult = {
            'Numkeyword': len(keyword),
            'CountLike': dem
        }
        # document = Document()
        endtime = datetime.now()
        end_time = endtime.strftime("%H:%M:%S")
        end_time = end_time.replace(":", "_")
        filename = None
        BotSessionResult = json.dumps(BotSessionResult)
        asyncio.create_task(post_log(filename, BotSessionResult))

    async def facebook_auto_unlike(group, keyword, from_time, to_time):
        print("Start unlike")
        dem = 0
        list_id = []
        for i in range(len(group)):
            data = get_posts(group=group[i]['id'],
                             cookie=r'C:\Users\Admin 3i\PycharmProjects\Source\facebook_scrapy\facebook_scrapy\spiders\cookie.json',
                             pages=2)
            data = pd.DataFrame.from_dict(data)
            a = 0
            while a < len(data):
                time_post = data['time'][a]
                content = data['text'][a]
                User = data['username'][a]
                print(User)
                if check_keyword(keyword, content) == 1:
                    if check_time(from_time, to_time, time_post) == 1:
                        list_id.append(data['post_id'][a])
                a += 1
            print(len(list_id))

            for product in list_id:
                async with websockets.connect(URL, ping_interval=None) as ws:
                    url = "https://www.facebook.com/" + str(product)
                    browser.get(url)
                    await ws.send("Unlike link: " + url)
                    wait()
                    try:
                        browser.find_element_by_xpath('//*[@aria-label="Gỡ Thích"]').click()
                        await ws.send("Done!")
                        dem += 1
                    except:
                        pass
        BotSessionResult = {
            'Numkeyword': len(keyword),
            'UnlikeCount': dem
        }
        filename = None
        BotSessionResult = json.dumps(BotSessionResult)
        asyncio.create_task(post_log(filename, BotSessionResult))

    async def facebook_auto_follow(name):
        print("Start follow!")
        for i in range(len(name)):
            print(name[i])
            async with websockets.connect(URL, ping_interval=None) as ws:

                browser.find_element_by_xpath('//*[@aria-label="Tìm kiếm trên Facebook"]').click()
                sleep(1)
                search = browser.find_element_by_xpath('//*[@aria-label="Tìm kiếm trên Facebook"]')
                search.send_keys(name[i])
                sleep(1)
                search.send_keys(Keys.ENTER)
                sleep(2)
                try:
                    sleep(2)
                    WebDriverWait(browser, 10).until(ec.element_to_be_clickable(
                        (By.XPATH, '//*[@aria-label="Theo dõi"]'))).click()
                    await ws.send("Followed: " + name[i])
                except:
                    await ws.send("Cant follow!")
                    pass

    async def facebook_auto_unfollow(name):
        for i in range(len(name)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                print("Start unfollow!")
                await ws.send("Start Unfollow!")
                try:
                    await ws.send("Search: " + name[i])
                    browser.find_element_by_xpath('//*[@aria-label="Tìm kiếm trên Facebook"]').click()
                    await asyncio.sleep(1)
                    search = browser.find_element_by_xpath('//*[@aria-label="Tìm kiếm trên Facebook"]')
                    search.send_keys(name[i])
                    await asyncio.sleep(1)
                    search.send_keys(Keys.ENTER)
                    sleep(2)
                    browser.find_element_by_xpath('//*[@aria-label="Bỏ theo dõi"]').click()
                    await ws.send("Unfollowed" + name[i])
                except:
                    pass
        browser.quit()

    async def fb_share_post(post, friend, group, page, content):
        for i in range(len(post)):
            url = 'https://www.facebook.com/' + post[i]['id']
            browser.get(url)
            sleep(3)
            # share to friend wall
            try:
                for j in range(len(friend)):
                    async with websockets.connect(URL, ping_interval=None) as ws:
                        await ws.send("Friend share!")
                        browser.find_element_by_xpath(
                            '//*[@aria-label="Gửi nội dung này đến bạn bè hoặc đăng trên dòng thời gian của bạn."]').click()
                        sleep(2)
                        browser.find_element_by_xpath(
                            '/html/body/div[1]/div/div[1]/div/div[3]/div/div/div[2]/div/div/div[1]/div[1]/div/div/div['
                            '1]/div/div/div[1]/div/div[7]/div/div[1]/div[2]/div/div/div/div/span/span').click()
                        sleep(2)
                        bd = browser.find_element_by_xpath(
                            '//*[@class="lzcic4wl gs1a9yip br7hx15l h2jyy9rg n3ddgdk9 owxd89k7 rq0escxv j83agx80 '
                            'a5nuqjux '
                            'l9j0dhe7 k4urcfbm rz7trki1 b3i9ofy5"]')
                        sleep(2)
                        bd.send_keys(friend[j])
                        sleep(1)

                        browser.find_element_by_xpath("//span[text()='" + friend[j] + "']").click()
                        sleep(2)
                        keyboard = Controller()
                        keyboard.type(content)
                        sleep(2)
                        await ws.send("Share to friend: " + friend[j])
                        browser.find_element_by_xpath('//*[@aria-label="Đăng"]').click()
                        sleep(2)

            except:
                pass
            # share to group
            try:
                for g in range(len(group)):
                    async with websockets.connect(URL, ping_interval=None) as ws:
                        await ws.send("Group share!")
                        try:
                            browser.find_element_by_xpath(
                                '//*[@aria-label="Gửi nội dung này đến bạn bè hoặc đăng trên dòng thời gian của bạn."]').click()
                            sleep(1)
                            browser.find_element_by_xpath(
                                '/html/body/div[1]/div/div[1]/div/div[3]/div/div/div[2]/div/div/div[1]/div['
                                '1]/div/div/div[1]/div/div/div[1]/div/div[5]/div/div[1]/div['
                                '2]/div/div/div/div/span/span').click()
                            sleep(1)
                            bd = browser.find_element_by_xpath(
                                '//*[@class="lzcic4wl gs1a9yip br7hx15l h2jyy9rg n3ddgdk9 owxd89k7 rq0escxv j83agx80 a5nuqjux l9j0dhe7 k4urcfbm rz7trki1 b3i9ofy5"]')

                            bd.send_keys(group[g]['id'])
                            sleep(1)
                            try:
                                browser.find_element_by_xpath(
                                    '/html/body/div[1]/div/div[1]/div/div[4]/div/div/div[1]/div/div[2]/div/div/div/div/div['
                                    '1]/div[2]/div[2]/div/div[2]/div/div[1]/div/div/div[2]/div/div/div/div[1]/div[2]/div['
                                    '1]').click()
                                keyboard = Controller()
                                keyboard.type(content)
                                browser.find_element_by_xpath('//*[@aria-label="Tạo bài viết công khai..."]').click()

                                sleep(2)
                                browser.find_element_by_xpath('//*[@aria-label="Đăng"]').click()
                                await ws.send("Share to group: " + group[g])
                            except:
                                pass
                        except:
                            await ws.send("Fail share group!")
                            pass
            except:
                pass
            try:
                for h in range(len(page)):
                    async with websockets.connect(URL, ping_interval=None) as ws:
                        await ws.send("Page share!")
                        try:
                            browser.find_element_by_xpath(
                                '//*[@aria-label="Gửi nội dung này đến bạn bè hoặc đăng trên dòng thời gian của bạn."]').click()
                            sleep(2)
                            # select option
                            browser.find_element_by_xpath(
                                '/html/body/div[1]/div/div[1]/div/div[3]/div/div/div[2]/div/div/div[1]/div['
                                '1]/div/div/div[1]/div/div/div[1]/div/div[6]/div/div[1]/div['
                                '2]/div/div/div/div/span/span').click()
                            sleep(1)
                            # click page
                            browser.find_element_by_xpath(
                                '/html/body/div[1]/div/div[1]/div/div[4]/div/div/div[1]/div/div[2]/div/div/div/div/div['
                                '1]/div[3]/div/div/div/div/div[2]/div/div[1]/div/div[1]').click()
                            sleep(2)
                            keyboard = Controller()
                            keyboard.type(content)
                            sleep(2)
                            browser.find_element_by_xpath('//*[@aria-label="Chia sẻ"]').click()
                            await ws.send("Shared to page: " + page[h]['id'])
                        except:
                            await ws.send("Fail share page")
                            pass
            except:
                pass

    # done invite group page
    async def auto_invite(group, page, friends):
        try:
            print("Start invite friends to page")
            await asyncio.sleep(3)
            for i in range(len(page)):
                # async with websockets.connect(URL, ping_interval=None) as ws:
                url = 'https://www.facebook.com/' + page[i]['id']
                browser.get(url)
                # await ws.send('Invite to page: ' + url)
                sleep(4)
                dem = 0
                for a in range(len(friends)):
                    browser.find_element_by_xpath('//*[@aria-label="Xem tất cả bạn bè"]').click()
                    sleep(2)
                    box = browser.find_element_by_xpath('//*[@aria-label="Search in All Friends"]')
                    box.clear()
                    box.send_keys(friends[a]['id'])
                    # await ws.send('Invite: ' + friends[a]['id'])
                    browser.find_element_by_xpath(
                        "//*[@class='d2edcug0 hpfvmrgz qv66sw1b c1et5uql lr9zc1uh a8c37x1j fe6kdd0r mau55g9w c8b282yb keod5gw0 nxhoafnm aigsh9s9 d3f4x2em iv3no6db a5q79mjw g1cxx5fr lrazzd5p knomaqxo']").click()
                    sleep(2)
                    try:
                        browser.find_element_by_xpath('//*[@role="checkbox"]').click()
                        browser.find_element_by_xpath(
                            '//*[@class="oajrlxb2 rq0escxv f1sip0of hidtqoto nhd2j8a9 datstx6m kvgmc6g5 cxmmr5t8 oygrvhab hcukyx3x b5wmifdl lzcic4wl jb3vyjys rz4wbd8a qt6c0cv9 a8nywdso pmk7jnqg j9ispegn kr520xx4 k4urcfbm"]').click()
                        browser.find_element_by_xpath('//*[@aria-label="Gửi lời mời"]').click()
                    except:
                        browser.find_element_by_xpath('//*[@class="cypi58rs pmk7jnqg fcg2cn6m tkr6xdv7"]').click()
                        # await ws.send('Intive: ' + friends[a]['id'])
                        pass
        except:
            pass
        sleep(2)
        try:
            print("Start invite friends to group")
            for i in range(len(group)):
                # async with websockets.connect(URL, ping_interval=None) as ws:
                url = 'https://www.facebook.com/groups/' + group[i]['id']
                browser.get(url)
                # await ws.send('Start with group: ' + url)
                dem = 0
                for a in range(len(friends)):
                    browser.find_element_by_xpath('//*[@aria-label="Mời"]').click()
                    box = browser.find_element_by_xpath('//*[@aria-label="Tìm bạn bè theo tên"]')
                    box.clear()
                    box.send_keys(friends[a]['id'])
                    sleep(2)
                    try:
                        browser.find_element_by_xpath('//*[@role="checkbox"]').click()
                        sleep(1)
                        browser.find_element_by_xpath('//*[@aria-label="Gửi lời mời"]').click()
                        dem += 1
                    except:
                        pass
                    browser.find_element_by_xpath('//*[@aria-label="Đóng"]').click()
                # await ws.send('Invite total: ' + str(dem))
        except:
            pass

    def value_to_float(x):
        if type(x) == float or type(x) == int:
            return x
        if 'K' in x:
            if len(x) > 1:
                return float(x.replace('K', '')) * 1000
            return 1000.0
        if 'M' in x:
            if len(x) > 1:
                return float(x.replace('M', '')) * 1000000
            return 1000000.0
        if 'B' in x:
            return float(x.replace('B', '')) * 1000000000
        else:
            return x
    async def facebook_auto_search_group(name):

            for i in range(len(name)):
                async with websockets.connect(URL, ping_interval=None) as ws:
                    await ws.send("Search: " + name[i])
                    string = name[i].replace(" ", "%20")
                    string = 'https://www.facebook.com/search/groups/?q=' + string
                    print(string)
                    await asyncio.sleep(7)
                    browser.get(string)
                    dem = 0
                    while dem < 10:
                        browser.find_element(By.XPATH, '//body').send_keys(Keys.END)
                        sleep(1)
                        dem += 1
                        # element = browser.find_element(By.XPATH,'//*[@span="x193iq5w xeuugli x13faqbe x1vvkbs x1xmvt09 x1lliihq x1s928wv xhkezso x1gmr53x x1cpjm7i x1fgarty x1943h6x xudqn12 x3x7a5m x6prxxf xvq8zen xo1l8bm xi81zsa x2b8uid"]')
                        # print(element)
                        # if element != None:
                        #     pass
                    page_source = browser.page_source
                    soup = BeautifulSoup(page_source, 'lxml')
                    # x9f619 x1n2onr6 x1ja2u2z x2bj2ny x1qpq9i9 xdney7k xu5ydu1 xt3gfkd xh8yej3 x6ikm8r x10wlt62 xquyuld
                    ListInfo = soup.find_all('div',
                                             class_='x9f619 x1n2onr6 x1ja2u2z x2bj2ny x1qpq9i9 xdney7k xu5ydu1 xt3gfkd xh8yej3 x6ikm8r x10wlt62 xquyuld')
                    count = 0
                    for info in ListInfo:
                        try:
                            count = count +1
                            joined = info.find('div', class_='x16n37ib x1n2onr6 x1e56ztr x1xmf6yo xamitd3').text
                            if 'Tham gia' in joined:
                                Isjoined = False
                            else:
                                Isjoined = True
                            element = info.find('span', class_='x193iq5w xeuugli x13faqbe x1vvkbs x1xmvt09 x1lliihq x1s928wv'
                                                               ' xhkezso x1gmr53x x1cpjm7i x1fgarty x1943h6x xudqn12 x676frb x1lkfr7t x1lbecb7 xk50ysn xzsf02u x1yc453h')

                            Grouppath = element.find('a')
                            GroupName = Grouppath.text
                            Grouppath = Grouppath['href']
                            GroupId = Grouppath.replace('https://www.facebook.com/groups/', "")
                            GroupId = GroupId.replace('/', "")
                            Frequecy = info.find('span',
                                                 class_='x193iq5w xeuugli x13faqbe x1vvkbs x1xmvt09 x1lliihq x1s928wv xhkezso x1gmr53x x1cpjm7i x1fgarty x1943h6x xudqn12 x3x7a5m x6prxxf xvq8zen xo1l8bm xi81zsa x1yc453h')

                            Frequecy = Frequecy.text
                            Permission = Frequecy.split('·')[0]
                            if 'Công khai' in Permission:
                                IsPermission = False
                            else:
                                IsPermission = True
                            MemberCount = Frequecy.split('·')[1]
                            MemberCount = MemberCount.replace("thành viên", '')
                            MemberCount = MemberCount.replace(" ", '')
                            MemberCount = MemberCount.replace(",", '.')
                            MemberCount = value_to_float(MemberCount)
                            FrequencyPost = None
                            FrequencyType = None
                            FrequencyPost = re.findall('(\d+)', (Frequecy.split('·')[2]))[0]
                            FrequencyType = Frequecy.split('·')[2].replace(FrequencyPost, "")
                        except:
                            pass
                        try:

                            Summary = info.find('span',
                                                class_='x193iq5w xeuugli x13faqbe x1vvkbs x1xmvt09 x1lliihq x1s928wv xhkezso x1gmr53x x1cpjm7i x1fgarty x1943h6x'
                                                       ' x4zkp8e x3x7a5m x1nxh6w3 x1sibtaa xo1l8bm xi81zsa x1yc453h')

                            Summary = Summary.text
                        except:
                            Summary = ""
                            pass
                        try:

                            Mutual = info.find('div', class_='x78zum5 x1q0g3np x6s0dn4')
                            Mutual = Mutual.find('span',
                                                 class_='x193iq5w xeuugli x13faqbe x1vvkbs x1xmvt09 x1lliihq x1s928wv xhkezso x1gmr53x x1cpjm7i x1fgarty x1943h6x xudqn12 x3x7a5m x6prxxf xvq8zen xo1l8bm xi81zsa').text
                            Mutual = re.findall('(\d+)', Mutual)[0]
                        except:
                            Mutual = None
                            pass
                        if FrequencyPost!= None:
                            FrequencyPost = int(FrequencyPost)
                        if Mutual != None:
                            Mutual = int(Mutual)
                        GroupInfo = {
                            'GroupId': GroupId,
                            'GroupName': GroupName,
                            'MemberCount': int(MemberCount),
                            'FrequencyPost': FrequencyPost,
                            'FrequencyType': FrequencyType,
                            'TypeGroup': 'DOANH_NGHIEP',
                            'Summary': Summary,
                            'MutualFriend': Mutual,
                            'IsPermission': IsPermission,
                            'IsJoin': Isjoined,
                            'IsDeleted': False,
                        }
                        print(GroupInfo)
                        url_upload = "http://localhost:6002/PythonCrawler/InsertFacebookGroupInfo"
                        resp = requests.post(url_upload, data= GroupInfo)
                        if resp.ok:
                            print("Upload completed successfully!")
                            print(resp.text)
                        else:
                            print("Something went wrong!")

    async def facebook_auto_join_group(group):

        for i in range(len(group)):
            async with websockets.connect(URL, ping_interval=None) as ws:
                await ws.send(group[i]["label"])
                idgroup = group[i]["id"]
                print(group)
                print(idgroup)
                print(group[i]["label"])
                string = 'https://www.facebook.com/' + idgroup
                print(string)
                await asyncio.sleep(7)
                browser.get(string)
                try:
                    # tham gia nhóm
                    browser.find_element(By.XPATH,
                                             '//*[@aria-label="Tham gia nhóm"]').click()
                    print('click')
                except:
                    pass
                try:
                    # theo dõi page
                    browser.find_element(By.XPATH,
                                             '//*[@aria-label="Theo dõi"]').click()
                    print('click')
                except:
                    pass
                #pass qua caai hỏi.
                try:
                    sleep(2)
                    WebDriverWait(browser, 10).until(ec.visibility_of_element_located(
                        (By.XPATH, '//*[@class="cypi58rs pmk7jnqg fcg2cn6m tkr6xdv7"]')))
                    try:
                        browser.find_element(By.XPATH,
                                             '//*[@class="l9j0dhe7 eg9m0zos m8pd07jg sjgh65i0 linmgsc8"]').send_keys(
                            Keys.END)
                    except:
                        pass
                    try:
                        browser.find_element(By.XPATH, '//*[@type="radio"]').click()
                    except:
                        pass

                    try:
                        boxtexts = browser.find_element(By.XPATH,
                                                        '//*[@class="oajrlxb2 f1sip0of hidtqoto lzcic4wl b3i9ofy5 l6v480f0 maa8sdkg s1tcr66n aypy0576 beltcj47 p86d2i9g aot14ch1 kzx2olss rq0escxv oo9gr5id l94mrbxd ekzkrbhg k4urcfbm o8yuz56k duhwxc4d bs68lrl8 f56r29tw e16z4an2 ei4baabg b4hei51z ehryuci6 hzawbc8m cxgpxx05 d1544ag0 sj5x9vvc tw6a2znq ieid39z1"]')
                        print(boxtexts)
                        for boxtext in boxtexts:
                            boxtext.send_keys("oke")
                            sleep(2)
                    except:
                        pass
                    try:
                        browser.find_element(By.XPATH,
                                             '//*[@class="rq0escxv du4w35lb l9j0dhe7 lzcic4wl j83agx80"]').click()
                        sleep(2)
                        browser.find_element(By.XPATH, '//*[@aria-label="Gửi"]').click()
                        await ws.send('Done!')
                        browser.find_element(By.XPATH,
                                             '//*[@class="l9j0dhe7 du4w35lb j83agx80 pfnyh3mw taijpn5t bp9cbjyn owycx6da btwxx1t3 kt9q3ron ak7q8e6j isp2s0ed ri5dt5u2 rt8b4zig n8ej3o3l agehan2d sk4xxmp2 rq0escxv tkv8g59h fl8dtwsd s1i5eluu tv7at329"]').click()
                    except:
                        pass
                    sleep(2)
                except:
                    pass

    async def post_log(filename, result):
        endtime = datetime.now()
        end_time = endtime.strftime("%H:%M:%S")
        end_time = end_time.replace(":", "_")
        data = {
            'SessionCode': Botname[0] + str(end_time),
            'StartTime': start_time,
            'EndTime': datetime.now(),
            'Statvs': '',
            'BotSocialCode': Botid[0],
            'FileResults': filename,
            'RuningType': 'facebook',
            'BotSessionResult': result,
            'CreatedBy': 'admin',
        }
        url_upload = "http://localhost:6023/PythonCrawler/InsertBotSocialLog"
        resp = requests.post(url_upload, data=data)
        if resp.ok:
            print("Upload completed successfully!")
            print(resp.text)
        else:
            print("Something went wrong!")

    def get_friend():
        # get friend for .NET data
        browser.find_element_by_xpath(
            '//*[@class="bp9cbjyn j83agx80 datstx6m taijpn5t oi9244e8 d74ut37n dt6l4hlj aferqb4h q5xnexhs"]').click()
        sleep(3)
        dem = 0
        while dem < 3:
            print('a')
            browser.find_element_by_xpath('//body').send_keys(Keys.END)
            sleep(1)
            dem += 1

        wait = WebDriverWait(browser, 10)
        wait.until(ec.visibility_of_element_located((By.LINK_TEXT, 'Xem tất cả bạn bè'))).click()
        sleep(5)

        while dem < 10:
            print('a')
            browser.find_element_by_xpath('//body').send_keys(Keys.END)
            sleep(1)
            dem += 1

        page_source = browser.page_source
        soup = BeautifulSoup(page_source, 'lxml')
        friend_local = soup.find('div', class_="j83agx80 btwxx1t3 lhclo0ds i1fnvgqd")
        friend_lst = friend_local.find_all('div',
                                           class_="bp9cbjyn ue3kfks5 pw54ja7n uo3d90p7 l82x9zwi n1f8r23x rq0escxv "
                                                  "j83agx80 bi6gxh9e discj3wi hv4rvrfc ihqw7lf3 dati1w0a gfomwglr")
        list_data = []
        for i in friend_lst:
            friend_link = i.find('div', class_="buofh1pr hv4rvrfc")
            # url = friend_link.find('a')['href']
            name = friend_link.find('span',
                                    class_="d2edcug0 hpfvmrgz qv66sw1b c1et5uql lr9zc1uh a8c37x1j fe6kdd0r mau55g9w "
                                           "c8b282yb keod5gw0 nxhoafnm aigsh9s9 d3f4x2em mdeji52x a5q79mjw g1cxx5fr "
                                           "lrazzd5p oo9gr5id").text
            print(name)
            data = {
                'id': name,
                'label': name
            }
            list_data.append(data)

        print(list_data)
        print(len(list_data))
        data_string = json.dumps(list_data, indent=4)
        with open('FacebookFriend.json', 'w') as f:
            f.write(data_string)

    async def add_friend(group):
        lst = []
        for i in range(len(group)):

            link = "https://www.facebook.com/groups/" + str(group[i]['id'])

            browser.get(link)
            await asyncio.sleep(2)
            try:
                browser.find_element_by_link_text("Mọi người").click()
                print(1)
            except:
                pass
            sleep(2)
            # kéo tới cuối trang để load tất cả các group
            reached_page_end = False
            last_height = browser.execute_script("return document.body.scrollHeight")
            # last_height = browser.execute_script("return document.documentElement.scrollHeight")
            print(last_height)
            while not reached_page_end:
                browser.find_element_by_xpath('//body').send_keys(Keys.END)
                sleep(2)
                new_height = browser.execute_script("return document.body.scrollHeight")
                print(new_height)
                if last_height == new_height:
                    reached_page_end = True
                else:
                    last_height = new_height
            sleep(5)

            page_source = browser.page_source
            soup = BeautifulSoup(page_source, "lxml")
            local = soup.find("div",
                              class_="rq0escxv l9j0dhe7 du4w35lb hpfvmrgz g5gj957u aov4n071 oi9244e8 bi6gxh9e h676nmdw aghb5jc5 gile2uim pwa15fzy fhuww2h9")
            tag = local.find_all('span', class_="nc684nl6")
            await asyncio.sleep(2)
            for j in tag:
                try:
                    link = j.find('a')['href']
                except:
                    pass
                print(link)
                lstnumb = re.findall("\d+", link)
                id = lstnumb[1]
                lst.append(id)

        df = pd.DataFrame(lst)
        df.to_csv(r'C:\Users\Admin 3i\PycharmProjects\Source\facebook_scrapy\logs/' + 'Memberid.csv', index=False,
                  header=False)
        sleep(2)
        products_data = pd.read_csv(
            r'C:\Users\Admin\PycharmProjects\Source\facebook_scrapy\logs\Memberid.csv')
        to_scan = products_data.loc[products_data['scanstatus'] != 'Yes']
        list_link = tqdm(to_scan['ID'].tolist())
        for id in list_link:
            link = "https://www.facebook.com/groups/" + str(id)
            browser.get(link)
            sleep(4)
            browser.find_element_by_xpath('//*[@aria-label="Thêm bạn bè"]').click()
            products_data.loc[products_data['ID'] == id, 'scanstatus'] = 'Yes'
            products_data.to_csv(r"C:\Users\Admin 3i\PycharmProjects\Source\facebook_scrapy\logs\Memberid.csv")

    async def get_comment_post(friend, keyword):
        async with websockets.connect(URL, ping_interval=None) as ws:
            print("Start follow!")
            await ws.send("Get comment post")
        for i in range(len(friend)):
            browser.get(friend)
            dem = 0
            while dem < 5:
                browser.find_element_by_xpath('//body').send_keys(Keys.END)
                sleep(2)
                dem += 1
            await asyncio.sleep(5)
            link_local = browser.find_elements_by_xpath(
                '//*[@aria-label="Gửi nội dung này đến bạn bè hoặc đăng trên dòng thời gian của bạn."]')
            print(link_local)  # list
            while dem < 20:
                sleep(1)
                browser.find_element_by_xpath('//body').send_keys(Keys.PAGE_UP)
                dem += 1
            LST_link = []
            for i in link_local:
                await asyncio.sleep(2)
                # wait = WebDriverWait(browser, 10)
                i.click()

                await asyncio.sleep(2)
                browser.find_element_by_xpath('//*[@rel="nofollow"]').click()
                await asyncio.sleep(2)
                browser.switch_to.window(browser.window_handles[1])
                page_source = browser.page_source
                soup = BeautifulSoup(page_source, 'lxml')
                try:
                    link = soup.find('div', class_="_2yyk _1hg8 _85ke").text
                except:
                    pass
                print(link)
                browser.close()
                browser.switch_to.window(browser.window_handles[0])
                LST_link.append(link)
            print(LST_link)
            Lst_content = []
            for j in LST_link:
                time_check = '1/1/2022'
                timeask = re.findall("\d+", time_check)
                monthcheck = int(timeask[1])
                daycheck = int(timeask[0])
                yearcheck = int(timeask[2])
                print(monthcheck, daycheck, yearcheck)
                browser.get(j)
                dem = 0
                while dem < 3:
                    browser.find_element_by_xpath('//body').send_keys(Keys.END)
                    await asyncio.sleep(2)
                    dem += 1
                page_source = browser.page_source
                soup = BeautifulSoup(page_source, 'lxml')
                try:
                    time = soup.find('b', class_="t5a262vz nc684nl6 ihxqhq3m l94mrbxd aenfhxwr l9j0dhe7 sdhka5h4").text
                except:
                    time = soup.find('span', class_="j1lvzwm4 stjgntxs ni8dbmo4 q9uorilb gpro0wi8").text
                if 'giờ' in time:
                    year = int(yearcheck)
                    month = int(monthcheck)
                    day = int(daycheck)
                elif 'lúc' in time and 'tháng' not in time:
                    if 'Hôm qua' in time:
                        year = int(yearcheck)
                        month = int(datetime.now().strftime("%m"))
                        day = int(datetime.now().strftime("%d")) - 1

                    else:
                        year = int(yearcheck)
                        time_now = re.findall("\d+", time)
                        month = int(time_now[1])
                        day = int(time_now[0])
                elif 'tháng' in time and ',' not in time:
                    timeask1 = re.findall("\d+", time)
                    print(timeask1)
                    month = int(timeask1[1])
                    day = int(timeask1[0])
                    year = int(yearcheck)
                else:
                    timeask1 = re.findall("\d+", time)
                    print(timeask1)
                    month = int(timeask1[1])
                    day = int(timeask1[0])
                    year = int(timeask1[2])

                print(month, day, year)

                if yearcheck > year:
                    continue
                else:  #
                    if monthcheck > month:
                        continue
                    else:
                        if daycheck > day:
                            continue
                        else:
                            pass
                time = str(day) + "-" + str(month) + '-' + str(year)
                Lst_content.append(time)
                try:
                    name = soup.find('div', class_="dati1w0a ihqw7lf3 hv4rvrfc ecm0bbzt").text
                except:
                    pass
                try:
                    content = soup.find('div',
                                        class_="bp9cbjyn j83agx80 cbu4d94t datstx6m taijpn5t pmk7jnqg j9ispegn kr520xx4 "
                                               "k4urcfbm").text
                except:
                    pass
                for x in keyword:
                    if x in content:
                        Lst_content.append(content)

                print(time)
            print(Lst_content)
            document = Document()
            for value in Lst_content:
                document.add_paragraph(value)
            document.save(r'C:\Users\Admin 3i\PycharmProjects\Source\facebook_scrapy\PostContent.docx')

    def facebook_auto_search_hashtag():
        hashtag = "#elonmusk"
        wait()
        try:
            browser.find_element_by_xpath('//*[@aria-label="Tìm kiếm trên Facebook"]').click()
            search = browser.find_element_by_xpath('//*[@aria-label="Tìm kiếm trên Facebook"]')
            search.send_keys(hashtag)
            wait()
            search.send_keys(Keys.ENTER)
        except:
            wait()
        sleep(1)

    def check_keyword(keyword, content):
        try:
            for x in keyword:
                if x in content:
                    return 1
            return 0
        except:
            pass

    def check_time(from_time, to_time, time_post):
        from_time1 = from_time.replace('T', ' ')
        from_time2 = from_time1.replace('.000Z', '')
        to_time1 = to_time.replace('T', ' ')
        to_time2 = to_time1.replace('.000Z', '')
        if datetime.strptime(str(time_post), "%Y-%m-%d %H:%M:%S") > datetime.strptime(from_time2,
                                                                                      "%Y-%m-%d %H:%M:%S") and datetime.strptime(
            str(time_post), "%Y-%m-%d %H:%M:%S") < datetime.strptime(to_time2, "%Y-%m-%d %H:%M:%S"):
            return 1
        else:
            return 0

    def check_friend(username, friend):
        for x in friend:
            if username == x:
                return 1
        return 0

    async def listen():
        ws_connect = websockets.connect('ws://127.0.0.1:9091', ping_interval=None)
        async with ws_connect as wb:
            await wb.send('Spider running!')

            while True:

                param = await wb.recv()
                if "UserName" in param:
                    data = json.loads(param)
                    print(type(data))
                    # tk = data['UserName']
                    # mk = data['Password']
                    Botname.append(data['BotSocialName'])
                    Botid.append(data['BotSocialCode'])
                    cookie = data['Cookie']
                    #print(cookie)
                    await asyncio.sleep(5)
                    asyncio.create_task(loginFacebookByCookie(cookie))

                if "Start post content" in param:
                    data = json.loads(param)
                    print(data)
                    linkpost_group1 = data["Start post content"]['Group']
                    #print(linkpost_group1)
                    post_content = data["Start post content"]['content']
                    #print(post_content)
                    await asyncio.sleep(5)
                    asyncio.create_task(facebook_auto_post_group(linkpost_group1, post_content))

                if "Start post comment" in param:
                    data = json.loads(param)
                    await asyncio.sleep(5)
                    group = data["Start post comment"]['Group']
                    content = data["Start post comment"]['content']
                    # friend = data["Start post comment"]['friends']
                    keyword = data["Start post comment"]['keywords']
                    from_time = data["Start post comment"]['from']
                    to_time = data["Start post comment"]['to']
                    asyncio.create_task(facebook_auto_comment(group, content, keyword, from_time, to_time))
                if "Start get content" in param:
                    data = json.loads(param)
                    group = data["Start get content"]['Group']
                    friend = data["Start get content"]['friends']
                    keyword = data["Start get content"]['keywords']
                    from_time = data["Start get content"]['from']
                    to_time = data["Start get content"]['to']
                    await asyncio.sleep(5)
                    asyncio.create_task(get_post_content(group, keyword, from_time, to_time))
                # if "Start get comment" in param:
                #     data = json.loads(param)
                #     group = data['Get']['Comment']['Group']
                #     friend = data['Get']['Comment']['friends']
                #     keyword = data['Get']['Comment']['keywords']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(get_comment_post(friend, keyword))
                # if "Like post" in param:
                #     data = json.loads(param)
                #     print(data)
                #     group = data['Like post']['Group']
                #     friend = data['Like post']['friends']
                #     keyword = data['Like post']['keywords']
                #     from_time = data['Like post']['from']
                #     to_time = data['Like post']['to']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(facebook_auto_like(group, keyword, from_time, to_time))
                # if "Like comment" in param:
                #     group = data['Like']['Comment']['Group']
                #     post = data['Like']['Comment']['post']
                #     friend = data['Like']['Comment']['friends']
                #     keyword = data['Like']['Comment']['keywords']
                #     from_time = data['Like']['Comment']['from']
                #     to_time = data['Like']['Comment']['to']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(facebook_auto_like(group, keyword, from_time, to_time))
                # if "Unlike" in param:
                #     if "Unlike post" in param:
                #         data = json.loads(param)
                #         group = data['Unlike post']['Group']
                #         friend = data['Unlike post']['friends']
                #         keyword = data['Unlike post']['keywords']
                #         from_time = data['Unlike post']['from']
                #         to_time = data['Unlike post']['to']
                #         await asyncio.sleep(5)
                #         asyncio.create_task(facebook_auto_unlike(group, keyword, from_time, to_time))
                #     if "Unlike comment" in param:
                #         group = data['Unlike']['Comment']['Group']
                #         post = data['Unlike']['Comment']['post']
                #         friend = data['Unlike']['Comment']['friends']
                #         keyword = data['Unlike']['Comment']['keywords']
                #         from_time = data['Unlike']['Comment']['from']
                #         to_time = data['Unlike']['Comment']['to']
                #         await asyncio.sleep(5)
                #         asyncio.create_task(facebook_auto_unlike(group, keyword, from_time, to_time))

                # if "Follow" in param:
                #     data = json.loads(param)
                #     name = data['Follow']['Name']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(facebook_auto_unfollow(name))
                # if "Unfollow" in param:
                #     data = json.loads(param)
                #     name = data['Follow']['Name']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(facebook_auto_unfollow(name))
                if "SearchGroup" in param:
                    data = json.loads(param)
                    name = data['SearchGroup']['Name']
                    asyncio.create_task(facebook_auto_search_group(name))
                if "JoinGroup" in param:
                    data = json.loads(param)
                    Group = data['JoinGroup']['Group']
                    asyncio.create_task(facebook_auto_join_group(Group))
                # if "Invite" in param:
                #     data = json.loads(param)
                #     print(data)
                #     group = data['Invite']['Group']
                #     friend = data['Invite']['friends']
                #     page = data['Invite']['page']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(auto_invite(group, page, friend))

                # if "Share post" in param:
                #     data = json.loads(param)
                #     post = data["Share post"]['post']
                #     group = data["Share post"]['group']
                #     page = data["Share post"]['page']
                #     friends = data["Share post"]['friends']
                #     content = data["Share post"]['content']
                #     await asyncio.sleep(5)
                #     asyncio.create_task(fb_share_post(post, friends, group, page, content))
                if "Addfriend" in param:
                    data = json.loads(param)
                    group = data['Addfriend']['Group']
                    asyncio.create_task(add_friend(group))

                if 'Stop' in param:
                    browser.quit()
                    sys.exit()

    async def main():
        task_1 = asyncio.create_task(listen())
        await asyncio.sleep(0.250)
        await task_1


    async def forever():
        while True:
            await main()
    loop = asyncio.get_event_loop()
    loop.run_until_complete(forever())

class Myspider(scrapy.Spider):
    name = 'CoreFacebookSpider'
    some_attribute = "Yes|No"
    # http://localhost:9080/crawl.json?spider_name=Tailieu&url=https://tailieu.vn/
    start_urls = ['https://facebook.com/']
    print("hello world !")
    Bot_func()

# if __name__ == '__main__':
#     asyncio.run(main())
